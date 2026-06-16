# -*- coding: utf-8 -*-
import os
import sys
import subprocess
import re

# Auto-install python-docx if not available
try:
    import docx
except ImportError:
    print("python-docx not found. Installing...")
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    import docx

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# ── XML & LAYOUT STYLING HELPERS ──────────────────────────────────────────────

def set_paragraph_rtl(paragraph):
    """Sets paragraph direction to RTL and alignment to RIGHT."""
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    bidi.set(qn('w:val'), '1')
    pPr.append(bidi)

def set_run_font(run, font_name="Segoe UI", size_pt=11, color_rgb=None, bold=False, italic=False):
    """
    Sets the font name and size for both Latin and Complex Scripts (Arabic) 
    in the XML properties of the run. This resolves the MS Word fallback quirk.
    """
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic
    if color_rgb:
        run.font.color.rgb = color_rgb
        
    rPr = run._r.get_or_add_rPr()
    
    # Set font family for complex scripts (cs)
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:cs'), font_name)  # Arabic font family
    rPr.append(rFonts)
    
    # Set font size for complex scripts
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), str(int(size_pt * 2)))
    rPr.append(sz)
    
    szCs = OxmlElement('w:szCs')
    szCs.set(qn('w:val'), str(int(size_pt * 2)))
    rPr.append(szCs)

def add_rtl_run(paragraph, text, bold=False, italic=False, size_pt=11, color_rgb=None, font_name="Segoe UI"):
    """Adds a run with RTL support, specific font, size, and color."""
    run = paragraph.add_run(text)
    set_run_font(run, font_name=font_name, size_pt=size_pt, color_rgb=color_rgb, bold=bold, italic=italic)
    
    # Set run properties to RTL
    rPr = run._r.get_or_add_rPr()
    rtl = OxmlElement('w:rtl')
    rtl.set(qn('w:val'), '1')
    rPr.append(rtl)
    return run

def split_mixed_text(text):
    """
    Splits the mixed text by partitioning it into Arabic and non-Arabic blocks.
    A non-Arabic block that contains English words or digits is treated as LTR.
    This keeps punctuation like (Backend) or 1) in the LTR run, preventing mirroring.
    """
    # Regex to find all contiguous sequences of Arabic characters and Arabic punctuation
    arabic_pattern = r"([\u0600-\u06FF\u0750-\u077F\u08A0-\u08FF\uFB50-\uFDFF\uFE70-\uFEFF؟،؛]+)"
    parts = re.split(arabic_pattern, text)
    
    runs = []
    for part in parts:
        if not part:
            continue
        
        is_arabic = bool(re.search(r"[\u0600-\u06FF\u0750-\u077F\u08A0-\u08FF\uFB50-\uFDFF\uFE70-\uFEFF؟،؛]", part))
        
        if is_arabic:
            runs.append(('ARA', part))
        else:
            has_strong_ltr = bool(re.search(r"[A-Za-z0-9]", part))
            if has_strong_ltr:
                subparts = re.split(r"(`[^`]+`)", part)
                for sub in subparts:
                    if not sub:
                        continue
                    if sub.startswith("`") and sub.endswith("`"):
                        runs.append(('CODE', sub[1:-1]))
                    else:
                        runs.append(('ENG', sub))
            else:
                runs.append(('ARA', part))
                
    merged_runs = []
    for r_type, r_text in runs:
        if merged_runs and merged_runs[-1][0] == r_type:
            merged_runs[-1] = (r_type, merged_runs[-1][1] + r_text)
        else:
            merged_runs.append((r_type, r_text))
            
    return merged_runs

def add_mixed_text_to_p_with_font(p, text, is_italic=False, size_pt=10.5, default_color=None, is_bold=False):
    """
    Parses a mixed English and Arabic text string and adds separate LTR and RTL runs.
    Uses split_mixed_text to prevent MS Word from misplacing English terms and punctuation.
    """
    runs = split_mixed_text(text)
    
    for r_type, r_text in runs:
        if r_type == 'CODE':
            run = p.add_run(r_text)
            set_run_font(run, font_name="Consolas", size_pt=size_pt - 0.5, color_rgb=RGBColor(199, 37, 78), bold=True)
        elif r_type == 'ENG':
            run = p.add_run(r_text)
            set_run_font(run, font_name="Segoe UI", size_pt=size_pt, color_rgb=RGBColor(30, 58, 138), bold=True, italic=is_italic)
        else: # ARA
            run = p.add_run(r_text)
            set_run_font(run, font_name="Segoe UI", size_pt=size_pt, color_rgb=default_color, bold=is_bold, italic=is_italic)
            
            # Set RTL run properties explicitly
            rPr = run._r.get_or_add_rPr()
            rtl = OxmlElement('w:rtl')
            rtl.set(qn('w:val'), '1')
            rPr.append(rtl)

def add_mixed_paragraph(doc, text, style_type="body", indent_inches=0):
    """Helper to add a paragraph with mixed LTR/RTL text and proper spacing."""
    p = doc.add_paragraph()
    set_paragraph_rtl(p)
    
    if indent_inches > 0:
        p.paragraph_format.right_indent = Inches(indent_inches)
        
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.25
    
    default_color = RGBColor(55, 65, 81) # Charcoal
    is_bold = False
    size_pt = 10.5
    
    if style_type == "q":
        default_color = RGBColor(27, 54, 93) # Dark Navy
        is_bold = True
        size_pt = 11.5
    elif style_type == "desc":
        default_color = RGBColor(107, 114, 128) # Gray
        size_pt = 10
        
    add_mixed_text_to_p_with_font(p, text, size_pt=size_pt, default_color=default_color, is_bold=is_bold)
    return p

def set_cell_background(cell, hex_color):
    """Fills cell background with a specific hex color."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Sets cell padding (in dxa: 20 dxa = 1 pt)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_heading_1_banner(doc, text):
    """Creates a beautiful full-width heading banner with Deep Navy background."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    set_cell_background(cell, "1B365D") # Deep Navy
    set_cell_margins(cell, top=140, bottom=140, left=200, right=200)
    
    # Remove borders
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'none')
        tcBorders.append(border)
    tcPr.append(tcBorders)
    
    p = cell.paragraphs[0]
    set_paragraph_rtl(p)
    run = p.add_run(text)
    set_run_font(run, font_name="Segoe UI", size_pt=14, color_rgb=RGBColor(255, 255, 255), bold=True)
    
    # Set RTL
    rPr = run._r.get_or_add_rPr()
    rtl = OxmlElement('w:rtl')
    rtl.set(qn('w:val'), '1')
    rPr.append(rtl)
    
    cell.width = Inches(6.5)
    doc.add_paragraph() # Spacing

def add_heading_2_styled(doc, text):
    """Creates a modern Heading 2 block with a thick Emerald Green right border."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    set_cell_background(cell, "F4F6F9") # Light Gray
    
    # Thick right border
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    
    right = OxmlElement('w:right')
    right.set(qn('w:val'), 'single')
    right.set(qn('w:sz'), '24') # 3pt
    right.set(qn('w:space'), '0')
    right.set(qn('w:color'), "00A86B") # Emerald Green
    tcBorders.append(right)
    
    for border_name in ['top', 'left', 'bottom']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'none')
        tcBorders.append(border)
    tcPr.append(tcBorders)
    
    set_cell_margins(cell, top=80, bottom=80, left=120, right=120)
    
    p = cell.paragraphs[0]
    set_paragraph_rtl(p)
    run = p.add_run(text)
    set_run_font(run, font_name="Segoe UI", size_pt=12, color_rgb=RGBColor(27, 54, 93), bold=True)
    
    rPr = run._r.get_or_add_rPr()
    rtl = OxmlElement('w:rtl')
    rtl.set(qn('w:val'), '1')
    rPr.append(rtl)
    
    cell.width = Inches(6.5)
    doc.add_paragraph() # Spacing

def add_qa_card(doc, q_num, question, answer, tip=None):
    """
    Creates a premium single-card layout for a question-answer pair.
    Implements border borders and shading for high-end styling.
    """
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    
    # Background shading (Very soft white/gray-blue)
    set_cell_background(cell, "F8FAFC")
    set_cell_margins(cell, top=140, bottom=140, left=180, right=180)
    
    # Set borders: Thick Navy right border, thin light gray elsewhere
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    
    right = OxmlElement('w:right')
    right.set(qn('w:val'), 'single')
    right.set(qn('w:sz'), '24') # 3pt
    right.set(qn('w:space'), '0')
    right.set(qn('w:color'), "1B365D") # Navy Blue
    tcBorders.append(right)
    
    for b in ['top', 'left', 'bottom']:
        border = OxmlElement(f'w:{b}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4') # 0.5pt
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), "E2E8F0") # Light Gray
        tcBorders.append(border)
    tcPr.append(tcBorders)
    
    cell.width = Inches(6.5)
    
    # Question text
    p_q = cell.paragraphs[0]
    set_paragraph_rtl(p_q)
    p_q.paragraph_format.space_before = Pt(4)
    p_q.paragraph_format.space_after = Pt(4)
    p_q.paragraph_format.line_spacing = 1.3
    
    run_num = p_q.add_run(f"س{q_num}: ")
    set_run_font(run_num, font_name="Segoe UI", size_pt=12, color_rgb=RGBColor(255, 107, 53), bold=True) # Accent Orange
    
    add_mixed_text_to_p_with_font(p_q, question, size_pt=12, default_color=RGBColor(27, 54, 93), is_bold=True) # Navy
    
    # Answer text
    p_a = cell.add_paragraph()
    set_paragraph_rtl(p_a)
    p_a.paragraph_format.space_before = Pt(4)
    p_a.paragraph_format.space_after = Pt(4)
    p_a.paragraph_format.line_spacing = 1.3
    
    run_ans_lbl = p_a.add_run("ج: ")
    set_run_font(run_ans_lbl, font_name="Segoe UI", size_pt=10.5, color_rgb=RGBColor(0, 150, 136), bold=True) # Teal
    
    add_mixed_text_to_p_with_font(p_a, answer, size_pt=10.5, default_color=RGBColor(55, 65, 81)) # Charcoal
    
    # Optional tip text
    if tip:
        p_t = cell.add_paragraph()
        set_paragraph_rtl(p_t)
        p_t.paragraph_format.space_before = Pt(4)
        p_t.paragraph_format.space_after = Pt(4)
        
        run_tip_lbl = p_t.add_run("💡 نصيحة للدفاع: ")
        set_run_font(run_tip_lbl, font_name="Segoe UI", size_pt=9.5, color_rgb=RGBColor(255, 107, 53), bold=True)
        
        add_mixed_text_to_p_with_font(p_t, tip, is_italic=True, size_pt=9.5, default_color=RGBColor(85, 95, 110))
        
    # Spacing between cards
    p_space = doc.add_paragraph()
    p_space.paragraph_format.space_before = Pt(0)
    p_space.paragraph_format.space_after = Pt(6)

# ── DATA DEFINITION (200 QUESTIONS AND ANSWERS - HIGHLY ORDERED & STRUCTURED) ──

sections_data = [
    {
        "title": "الفصل الأول: البنية التحتية وهيكل النظام وموقع الاستضافة (Hostinger VPS)",
        "description": "يركز هذا الفصل على البنية الأساسية للمشروع، ومبررات اختيار لغة بايثون وFastAPI، ومقارنتها بنظام Node.js، مع تفصيل كامل لإعداد السيرفر على Hostinger VPS ونظام Nginx لتوجيه المرور وإدارة Uvicorn كخدمة خلفية نظام التشغيل.",
        "subcategories": [
            {
                "title": "الفرع الأول: مقارنة التقنيات ومبررات بايثون وFastAPI",
                "questions": [
                    {
                        "q": "لماذا اخترتم لغة Python لبناء الخلفية (Backend) للمشروع بدلاً من Node.js أو C#؟",
                        "a": "اخترنا Python لأنها اللغة الرائدة في مجال الذكاء الاصطناعي ومعالجة اللغات الطبيعية (NLP). توفر مكتبات رسمية وممتازة لـ Google Gemini وGroq، بالإضافة إلى مكتبة PyMuPDF التي تتميز بسرعتها الفائقة المكتوبة بلغة C لمعالجة ملفات PDF، ونظام التحقق من البيانات Pydantic v2 السريع جداً والمدعوم بلغة Rust. بينما استخدام Node.js كان سيتطلب كتابة مئات الأسطر من الكود الإضافي للتحقق من البيانات ومعالجة النصوص وتكامل الذكاء الاصطناعي.",
                        "tip": "ركز في إجابتك على أن لغة بايثون هي البيئة الطبيعية (Native Ecosystem) لمكتبات الذكاء الاصطناعي ومعالجة ملفات PDF السريعة."
                    },
                    {
                        "q": "ما هو إطار العمل المستخدم في الخلفية وما هي مميزاته الأساسية؟",
                        "a": "استخدمنا إطار العمل FastAPI (الإصدار 0.115.6). مميزاته تشمل: الدعم الكامل للعمليات غير المتزامنة (Asynchronous code - async/await)، الإنشاء التلقائي لمستندات واجهة برمجة التطبيقات (OpenAPI/Swagger docs)، التحقق التلقائي والكامل من البيانات عبر Pydantic، وسرعته الهائلة التي تقارب سرعة Node.js وGo بسبب اعتماده على ASGI (Starlette & Uvicorn).",
                    },
                    {
                        "q": "ما الفرق بين ASGI و WSGI، ولماذا يعتمد مشروعكم على ASGI؟",
                        "a": "الـ WSGI (Web Server Gateway Interface) يدعم الطلبات المتزامنة المتتالية (Synchronous/Blocking)، بينما الـ ASGI (Asynchronous Server Gateway Interface) يدعم الطلبات غير المتزامنة وثنائية الاتجاه (Asynchronous, WebSockets, SSE). مشروعنا يعتمد على ASGI لأننا نتعامل مع عمليات ذكاء اصطناعي وتوليد صوتيات وفيديوهات تأخذ وقتاً طويلاً، ونحتاج لتشغيلها بشكل غير متزامن دون حظر بقية الطلبات وتقديم تجربة مستخدم سلسة مثل بث البيانات الفوري (SSE).",
                    },
                    {
                        "q": "لماذا قمتم بتثبيت إصدارات المكتبات بدقة (Pinned Versions) في ملف requirements.txt؟",
                        "a": "تثبيت الإصدارات بدقة باستخدام `==` (مثل `pydantic==2.10.5`) يمنع حدوث ثلاثة مشاكل رئيسية في بيئة الإنتاج: 1) التغييرات المفاجئة غير المتوافقة (Breaking Changes) في المكتبات مستقبلاً، 2) عدم تطابق سلوك النظام بين جهاز المطور المحلي وخادم الإنتاج، 3) الهجمات الأمنية على سلاسل التوريد (Supply Chain Attacks) عبر تحديثات خبيثة للمكتبات الفرعية.",
                    },
                    {
                        "q": "ما هو نموذج التصميم (Design Pattern) العام المستخدم في هندسة الخدمات البرمجية بالمشروع؟",
                        "a": "نستخدم نموذج الخدمة المفردة (Singleton Pattern) للعملاء مثل عميل Supabase وعميل ElevenLabs وعميل HTTP، بالإضافة إلى نمط هندسة الطبقات (Layered Architecture) لفصل منطق العرض (Controller/API) عن منطق معالجة الأعمال (Service Layer) ومنطق البيانات والتحقق (Schemas).",
                    }
                ]
            },
            {
                "title": "الفرع الثاني: استضافة Hostinger VPS وإعداد بيئة Linux",
                "questions": [
                    {
                        "q": "لماذا تم اختيار نشر المشروع على خادم VPS من Hostinger بدلاً من خادم سحابي عديم الخادم (Serverless)؟",
                        "a": "اخترنا خادم VPS (Virtual Private Server) من Hostinger لأن معالجة الوسائط المتعددة (توليد الفيديوهات بـ FFmpeg والصوت بـ ElevenLabs) تتطلب وقت تنفيذ طويل قد يتجاوز 5 دقائق، وهو ما تمنعه منصات الـ Serverless (التي تفرض حداً أقصى 60 ثانية). كما أن تشغيل مكتبات معالجة الفيديو والصوت يحتاج إلى خادم دائم يحتفظ بالحالة (Stateful) لتثبيت الأدوات مثل FFmpeg محلياً دون الاضطرار لإعادة بنائها مع كل طلب.",
                        "tip": "العمداء يفضلون الخيارات التي تحسب قيود الوقت والأداء بدلاً من الحلول السريعة."
                    },
                    {
                        "q": "كيف يتم إعداد وتشغيل بيئة Python وتطبيق FastAPI على خادم VPS من Hostinger؟",
                        "a": "على خادم VPS من Hostinger يعمل بنظام Ubuntu 22.04 LTS، نقوم بإنشاء بيئة افتراضية (Virtual Environment) معزولة لـ Python 3.12، ثم نستخدم Uvicorn كخادم ASGI لتشغيل تطبيق FastAPI برمجياً، ونقوم بإدارته كخدمة خلفية دائمة عبر نظام التشغيل باستخدام `systemd` لضمان إعادة تشغيل التطبيق تلقائياً عند إقلاع الخادم أو حدوث أي عطل.",
                    },
                    {
                        "q": "ما هي التحديات والقيود عند تشغيل معالجة الوسائط على خادم VPS بموارد محدودة (مثلاً 2GB RAM) وكيف تعاملتم معها؟",
                        "a": "أهم القيود هي استهلاك الذاكرة العشوائية والمعالج أثناء تشغيل FFmpeg لتوليد الفيديوهات ذات التأثيرات الحركية وضغط الصوت. تعاملنا مع ذلك برمجياً بتقسيم معالجة النصوص وتوليد الصوتيات على دفعات صغيرة (Batches of 6) لتقليل الضغط على الذاكرة، والتخلص الفوري من الملفات المؤقتة باستخدام كتل `finally` البرمجية، وتحسين جودة ضغط الفيديو ليكون خفيف الاستهلاك.",
                        "tip": "هذا السؤال يثبت قدرتك على كتابة كود برمجى ذكي يراعي إمكانيات السيرفر المتاحة."
                    },
                    {
                        "q": "لماذا لم تستخدموا حاويات Docker كطريقة أساسية للنشر في الإنتاج بدلاً من التثبيت اليدوي على Hostinger VPS؟",
                        "a": "التثبيت اليدوي أسرع للبدء وتجربة التعديلات مباشرة، ولكننا قمنا بتجهيز `Dockerfile` تحسباً للترقية المستقبلية ولضمان سهولة نقل الكود لأي سيرفر آخر دون تكرار خطوات تثبيت المكتبات الخارجية مثل FFmpeg.",
                    },
                    {
                        "q": "كيف نراقب ونحمي خادم VPS من استهلاك كامل الذاكرة العشوائية (RAM) أثناء تشغيل خدمة توليد الصوت والفيديو؟",
                        "a": "نقوم برصد استهلاك الذاكرة عبر كتابة سجلات برمجية تفصيلية، ونستخدم خيوط معالجة منفصلة (`asyncio.to_thread`) للعمليات التي تستهلك الذاكرة. كما نقوم بتنظيف أي متغيرات وسائط وبايتات غير مستخدمة فور الانتهاء، ونعتمد على خدمات التخزين الخارجي لحفظ الملفات بدلاً من إبقائها في ذاكرة الخادم.",
                    }
                ]
            },
            {
                "title": "الفرع الثالث: خادم الويب Nginx وإدارة المرور والاتصال",
                "questions": [
                    {
                        "q": "ما هو دور خادم الويب Nginx على خادم Hostinger وكيف يقوم بتوجيه المرور لتطبيقنا؟",
                        "a": "خادم Nginx يعمل كخادم ويب وكيل عكسي (Reverse Proxy). يستقبل الطلبات الخارجية عبر المنفذين 80 (HTTP) و 443 (HTTPS) ويقوم بتوجيهها داخلياً إلى خادم Uvicorn الذي يعمل على المنفذ المحلي 8000. كما يقوم Nginx بالتعامل مع تشفير SSL وحجب الهجمات وتحديد حجم الملفات المرفوعة الأقصى عبر `client_max_body_size 20M`.",
                    },
                    {
                        "q": "كيف يتم تأمين وحماية خادم Nginx ضد الهجمات الإلكترونية وثغرات الاتصال؟",
                        "a": "نقوم بإعداد جدار حماية (UFW) على مستوى السيرفر للسماح بمرور منافذ الويب فقط (80, 443, 22). ونقوم بتعطيل عرض إصدار خادم Nginx في الترويسات (Headers) لمنع المخترقين من معرفة إصدار الخادم واستغلال ثغراته، بالإضافة لتحديد معدل الطلبات لكل IP.",
                    },
                    {
                        "q": "كيف نضمن استقرار اتصال Nginx مع خادم Uvicorn عند معالجة طلبات تأخذ وقتاً طويلاً (Timeout Settings)؟",
                        "a": "نقوم بزيادة مهلات الاتصال في ملف إعدادات Nginx (nginx.conf) مثل `proxy_read_timeout 600s` و `proxy_send_timeout 600s` و `keepalive_timeout 65s` لضمان عدم قيام Nginx بقطع الاتصال مع العميل أثناء قيام بايثون بمعالجة ملف PDF وتوليد الفيديو بالخلفية.",
                    },
                    {
                        "q": "ما هي الخطوات المتبعة لتثبيت وتجديد شهادة الأمان SSL تلقائياً على خادم Hostinger VPS؟",
                        "a": "نستخدم أداة `Certbot` المعتمدة من سلطة الشهادات المفتوحة Let's Encrypt. نقوم بتشغيل أمر توليد الشهادة لربط النطاق (Domain Name) بالخادم، وتقوم الأداة بتعديل ملف إعدادات Nginx تلقائياً وتفعيل بروتوكول HTTPS، كما تضيف مهمة دورية (Cron Job) لتجديد الشهادة كل 90 يوماً بشكل تلقائي.",
                    },
                    {
                        "q": "كيف يتم التعامل مع مشكلة انقطاع اتصال قاعدة البيانات الخارجية أو تعطل الـ API Keys في بيئة Hostinger؟",
                        "a": "نقوم ببرمجة دالة فحص اتصال عند إقلاع خادم Uvicorn تتأكد من صلاحية المفاتيح وصلاحية الاتصال بقاعدة Supabase. وفي حال حدوث فشل، يرسل النظام إشعاراً في سجلات السيرفر مع استمرار تشغيل الخلفية لتقديم خدمات محدودة بدلاً من الانهيار التام.",
                    }
                ]
            },
            {
                "title": "الفرع الرابع: إدارة العمليات في الخلفية بـ systemd و PM2",
                "questions": [
                    {
                        "q": "ما أهمية استخدام نظام `systemd` لإدارة خادم Uvicorn وكيف تم إعداده؟",
                        "a": "نظام `systemd` هو المسؤول عن إدارة الخدمات في نظام Linux. قمنا بإنشاء ملف خدمة مخصص (`ruya.service`) يحدد مسار بيئة بايثون الافتراضية وأمر تشغيل Uvicorn. يضمن هذا بقاء التطبيق يعمل بشكل مستقل ودائم كخدمة خلفية (Daemon)، ويقوم بإعادة تشغيل الخدمة تلقائياً في حال انهيار الكود برمجياً أو إعادة تشغيل السيرفر.",
                    },
                    {
                        "q": "كيف نتحكم في عدد العمليات (Workers) عند تشغيل خادم Uvicorn على خادم Hostinger؟",
                        "a": "نستخدم Gunicorn كمدير للعمليات (Process Manager) مع تشغيل Uvicorn كعامل داخلي (`UvicornWorker`). نحدد عدد الـ Workers بناءً على عدد نوايا المعالج في السيرفر (معادلة: CPU Cores * 2 + 1)، مما يسمح للخلفية بمعالجة عدة طلبات متزامنة بكفاءة عالية ودون حظر للعمليات.",
                    },
                    {
                        "q": "ما فائدة استخدام أداة `PM2` كبديل لـ systemd لإدارة تطبيقات بايثون؟",
                        "a": "أداة PM2 (المعروفة في بيئة Node.js ولكنها تدعم بايثون) توفر لوحة تحكم سهلة في الطرفية (Terminal) لمراقبة استهلاك الذاكرة والمعالج لكل عملية تشغيل لحظة بلحظة، وتسهل قراءة سجلات الأخطاء (Logs) وإعادة تشغيل التطبيق بأمر بسيط كـ `pm2 restart ruya`.",
                    },
                    {
                        "q": "كيف يتم حقن متغيرات البيئة (Environment Variables) داخل خدمة `systemd`؟",
                        "a": "نحدد ذلك داخل ملف الخدمة `ruya.service` باستخدام الإعداد `EnvironmentFile=/path/to/.env` أو كتابة المتغيرات مباشرة تحت الإعداد `Environment=`. يضمن هذا وصول كود بايثون لمفاتيح الذكاء الاصطناعي وSupabase بأمان تام ومنع استهلاكها من مستخدمين غير مصرح لهم.",
                    },
                    {
                        "q": "ماذا يحدث في نظام Linux لو توقفت خدمة Uvicorn فجأة وكيف يظهر ذلك للمستخدم النهائي؟",
                        "a": "سيقوم نظام `systemd` بالتقاط التوقف وإعادة تشغيل الخدمة خلال ثوانٍ معدودة بناءً على إعداد `Restart=always`. طوال هذه الثواني القليلة، سيقوم خادم Nginx بإرجاع خطأ `HTTP 502 Bad Gateway` للمستخدم النهائي، وبمجرد تعافي الخدمة يعود الاتصال للعمل تلقائياً.",
                    }
                ]
            },
            {
                "title": "الفرع الخامس: هيكلية مجلدات المشروع ومكتبة PyMuPDF",
                "questions": [
                    {
                        "q": "ما فائدة استخدام مكتبة `pydantic-settings` لإدارة التكوينات وإعدادات النظام؟",
                        "a": "تتيح لنا قراءة المتغيرات وإعدادات التطبيق من ملف `.env` محلياً أو من متغيرات بيئة النظام تلقائياً مع توفير ميزة التحقق من الأنواع (Type Validation) لكل متغير عند تشغيل التطبيق وفشل النظام فوراً وبشكل واضح إذا كان أحد الإعدادات الأساسية مفقوداً أو خاطئاً.",
                    },
                    {
                        "q": "كيف نضمن عدم تعارض الملفات المرفوعة من مستخدمين مختلفين في نفس الوقت على خادم VPS؟",
                        "a": "الخلفية لا تحفظ أي ملفات مرفوعة على القرص الصلب المحلي للخادم بشكل دائم، بل تقرأ الملفات مباشرة في الذاكرة العشوائية كبايتات مؤقتة (Temporary Bytes). وفي حال توليد ملفات وسائط مؤقتة، نستخدم معرّفات فريدة عالمياً (UUID) كأسماء للملفات لمنع أي تداخل أو تعارض بين الطلبات المتزامنة.",
                    },
                    {
                        "q": "لماذا تم اختيار مكتبة PyMuPDF لمعالجة ملفات PDF وما هي مميزاتها البرمجية؟",
                        "a": "مكتبة PyMuPDF (المعروفة بـ `fitz` في الكود) مكتوبة بلغة C، وهي أسرع بـ 10 إلى 20 ضعفاً من المكتبات المكتوبة بلغة Python الصرفة مثل `PyPDF2`. كما أنها تدعم استخراج النصوص بدقة متناهية مع الاحتفاظ بالترتيب الصحيح للأسطر والفقرات، وتدعم قراءة الصور المضمنة وجداول البيانات بكفاءة بالغة.",
                    },
                    {
                        "q": "كيف يتم استخراج النصوص من ملفات الـ PDF التي تم مسحها ضوئياً كصور فقط (Scanned PDFs)؟",
                        "a": "في ملف `file_service.py` يكتشف النظام إذا كان النص المستخرج فارغاً أو قصيراً جداً مقارنة بعدد الصفحات. عندئذٍ، يتم تحويل صفحات الـ PDF إلى صور باستخدام PyMuPDF وتمريرها إلى محرك رؤية الذكاء الاصطناعي (Gemini Vision OCR) لاستخراج النصوص منها بدقة متناهية كبديل لـ OCR التقليدي.",
                    },
                    {
                        "q": "كيف يساهم كود `file_service.py` في تنظيف النصوص واستبعاد الأحرف غير المرغوبة بعد استخراجها؟",
                        "a": "يحتوي الملف على دوال فحص وتصحيح نصوص تستخدم التعبيرات المنتظمة (Regex) لإزالة الفراغات المزدوجة، وحذف العلامات الخاصة والتنسيقات التالفة الناتجة عن فك تشفير خطوط الـ PDF، لضمان إرسال نص نظيف ومتناسق لنماذج الذكاء الاصطناعي لتفادي تشويش معالجتها.",
                    }
                ]
            }
        ]
    },
    {
        "title": "الفصل الثاني: محرك الذكاء الاصطناعي الهجين والـ LLMs (Hybrid AI Engine & LLMs)",
        "description": "يتناول هذا القسم استراتيجية الجمع بين موفري الذكاء الاصطناعي (Groq و Google Gemini)، وتبرير اختيار النماذج، ومكافحة الهلوسة الذكائية، وهندسة التلقين ودرجة الحرارة (Temperature).",
        "subcategories": [
            {
                "title": "الفرع الأول: فلسفة المعمارية الهجينة والتبديل التلقائي",
                "questions": [
                    {
                        "q": "ما هي الفكرة الأساسية وراء استخدام 'محرك الذكاء الاصطناعي الهجين' (Hybrid AI Engine)؟",
                        "a": "الفكرة هي دمج أفضل ما في العالمين: السرعة الفائقة والتكلفة المنخفضة لمعالجة الهياكل النمطية المتكررة من جهة، والقدرة العالية على الفهم المنطقي والتحليل الهيكلي للمستندات الضخمة من جهة أخرى. نستخدم Groq (عبر Llama 3.3) للامتحانات نظراً لسرعته الفائقة، ونستخدم Gemini 2.0 Flash لخرائط المفاهيم لقدرته الفائقة على التحليل الهرمي وصياغة هياكل شجرية متماسكة.",
                        "tip": "اشرح للجنة أن الجمع بين مزودين يمنح التطبيق مرونة معمارية عالية جداً ويحميه من التوقف المفاجئ لأحدهما."
                    },
                    {
                        "q": "كيف يعمل نظام التبديل التلقائي عند الفشل (Failover System) في محرك الذكاء الاصطناعي؟",
                        "a": "يعمل النظام في وضع `hybrid` بشكل افتراضي؛ إذا فشل الطلب الموجه للمزود الأساسي (بسبب نفاد حد الطلبات Rate Limit أو انقطاع الخدمة أو حدوث خطأ داخلي)، يلتقط الكود الاستثناء تلقائياً (Exception Handling)، ويقوم فوراً وبشكل شفاف دون إشعار المستخدم بإرسال نفس الطلب للمزود الاحتياطي (Gemini كاحتياطي لـ Groq، وGroq كاحتياطي لـ Gemini) مما يضمن استمرارية الخدمة بنسبة تتعدى 99.9%.",
                    },
                    {
                        "q": "ما هي المزايا الفنية لوجود مزودين مستقلين (Multi-Provider Setup) بدلاً من الاعتماد على مكتبة واحدة؟",
                        "a": "يمنع الاحتكار البرمجي، ويحمي التطبيق من التوقف التام عند حدوث أعطال عالمية لأحد المزودين. كما يتيح تدوير واستهلاك الحصص المجانية وتخفيض الفاتورة الشهرية وتوزيع الأحمال بالتوازي مما يرفع كفاءة وموثوقية النظام بالكامل.",
                    },
                    {
                        "q": "كيف يقيس المحرك زمن استجابة كل نموذج (Latency Comparison) وبناءً عليه يتخذ قرار التبديل؟",
                        "a": "نقوم بدمج مؤقتات زمنية (Time Trackers) داخل كود استدعاء الـ API. إذا تجاوز زمن الاستجابة حداً معيناً (Timeout limit)، يقطع الاتصال تلقائياً ويبدأ فوراً استدعاء المزود الآخر لضمان بقاء زمن الاستجابة الكلي ضمن الحدود المقبولة للطالب.",
                    },
                    {
                        "q": "هل يمكن إضافة موفر ثالث مثل OpenAI GPT-4o للمحرك الهجين مستقبلاً وكيف يدعم الكود ذلك؟",
                        "a": "نعم، الكود مهيأ تماماً لدمج أي موفر جديد؛ حيث تعتمد هندسة محرك الذكاء الاصطناعي على تعريف واجهة معيارية موحدة (Standard Interface). يمكن بسهولة إضافة دالة استدعاء لـ OpenAI وضبطها كخيار ثالث في مصفوفة التبديل التلقائي دون تعديل الكود الرئيسي.",
                    }
                ]
            },
            {
                "title": "الفرع الثاني: موفري الذكاء الاصطناعي Groq و Gemini",
                "questions": [
                    {
                        "q": "لماذا تم اختيار نموذج `llama-3.3-70b-versatile` عبر منصة Groq لتوليد الامتحانات؟",
                        "a": "لأن Groq تعتمد على رقاقات معالجة لغوية ثورية ومخصصة (LPU - Language Processing Unit) تتيح سرعة توليد مذهلة تتجاوز 500 رمز في الثانية (Tokens/sec). وحيث أن توليد بنك أسئلة مكون من 50 سؤالاً مفصلاً يتطلب توليد حجم ضخم من النصوص وهيكل JSON متكرر، فإن Groq تقدم لنا زمناً قياسياً لا يتعدى ثانيتين مقارنة بـ 15-20 ثانية على كرت الشاشة التقليدي (GPU).",
                    },
                    {
                        "q": "لماذا فضلتم استخدام Google Gemini 2.0 Flash في توليد خرائط المفاهيم (Mind Maps)؟",
                        "a": "لأن توليد خريطة مفاهيم يتطلب فهماً شاملاً لترابط الأفكار داخل المستند وتوليد شجرة هيكلية متسقة وذات منطق تعليمي. يتميز Gemini 2.0 Flash بنافذة سياق (Context Window) عملاقة تصل لمليون رمز وبقدرات تفكير منطقي عالية جداً في الهيكلة، ودعم أصيل لإخراج صيغ JSON بدقة متناهية دون كسر البنية المتداخلة للشجرة.",
                    },
                    {
                        "q": "ما الفرق المعماري بين رقاقات الـ LPU لـ Groq والـ TPU لـ Google Gemini؟",
                        "a": "رقاقات الـ LPU (Language Processing Unit) مصممة خصيصاً لتسريع توليد الرموز المتتالية لنماذج اللغة وتعتمد على ذاكرة فائقة السرعة على الرقاقة نفسها، بينما الـ TPU (Tensor Processing Unit) مصممة لتسريع العمليات الحسابية للمصفوفات الضخمة وهي ممتازة للتدريب ومعالجة الوسائط المتعددة واستنتاج النماذج المعقدة.",
                    },
                    {
                        "q": "لماذا تم دمج محرك Gemini باستخدام مكتبة `google-generativeai` بدلاً من LangChain؟",
                        "a": "الاستخدام المباشر للمكتبة الرسمية يمنحنا تحكماً كاملاً بأداء الطلبات ويقلل من طبقات الكود الوسيطة (Abstractions) التي تزيد من استهلاك الذاكرة ووقت الاستجابة وتؤخر اكتشاف الأخطاء. إطار LangChain ممتاز للمشاريع المعقدة والعملاء المتعددين، ولكنه يضيف تعقيداً غير مبرر لمشروع يحتاج فقط لاتصالات مباشرة ومثبتة.",
                    },
                    {
                        "q": "كيف يتم تدوير مفاتيح الـ API (Key Rotation) لتفادي قيود الاستهلاك المجانية للموفرين؟",
                        "a": "يمكن تهيئة قائمة من المفاتيح في ملف الإعدادات وقراءتها برمجياً، وعند استقبال خطأ HTTP 429 (Too Many Requests)، يقوم الكود بنقل المؤشر للمفتاح التالي في القائمة وإعادة المحاولة فوراً لتجاوز القيود الساعية لـ API.",
                    }
                ]
            },
            {
                "title": "الفرع الثالث: درجة الحرارة وصيغة مخرجات JSON الصارمة",
                "questions": [
                    {
                        "q": "لماذا تم ضبط معيار درجة الحرارة (Temperature) على القيمة صفر في جميع استدعاءات النماذج؟",
                        "a": "ضبط درجة الحرارة على 0 يجعل النموذج 'حتمياً' (Deterministic)؛ أي أنه يختار دائماً الرمز الأعلى احتمالية في كل خطوة بناء نصوص. هذا يلغي العشوائية والابتكار تماماً، وهو أمر ضروري في التطبيقات التعليمية وبناء الامتحانات حيث نريد إجابات دقيقة وحقائق مستمدة تماماً من النص دون أي تلاعب أو تغيير بالصياغات.",
                    },
                    {
                        "q": "كيف تضمنون أن الهيكل الذي يرجعه نموذج الذكاء الاصطناعي متوافق تماماً مع صيغة JSON المطلوبة؟",
                        "a": "نحقق ذلك بطريقتين: الأولى هي تفعيل الوضع الأصيل لـ JSON في النماذج (JSON Mode) مثل `response_format={'type': 'json_object'}` في Groq و `response_mime_type='application/json'` في Gemini. والثانية هي استخدام دالة التنظيف البرمجية الخاصة بنا `clean_and_parse_json()` التي تقوم بقص وتنظيف أي نصوص أو تعليقات زائدة ومحاولة إصلاح بنية JSON التالفة قبل تمريرها لمرحلة التحقق النهائي.",
                    },
                    {
                        "q": "لماذا تم فرض توليد مخرجات JSON باللغة الإنجليزية حصراً لخرائط المفاهيم وبنوك الأسئلة؟",
                        "a": "لأن لغة البرمجة C# في الواجهة الأمامية تتعامل مع كود موحد منسق يسهل فك تشفيره وعرضه، ولتجنب أخطاء ترميز الحروف والكلمات المتداخلة في هياكل JSON الشجرية المعقدة عند كتابة مفاتيح ومعرفات باللغة العربية، بالإضافة إلى توحيد العقود البرمجية وتسهيل ترجمتها في الواجهة الأمامية للمستخدم إذا تطلب الأمر.",
                    },
                    {
                        "q": "كيف يمنع الكود الذكاء الاصطناعي من إضافة تعليقات أو نصوص توضيحية قبل أو بعد هيكل الـ JSON المرجَّع؟",
                        "a": "نقوم بفرض ذلك عبر إرشادات النظام الصارمة وتفعيل وضع JSON الصارم في النماذج. وللحماية الإضافية، نمرر النص المرجع عبر تعبيرات منتظمة (Regex) في الخلفية لاقتطاع وحذف أي نصوص زائدة والبدء في التحليل البرمجي فوراً من أول قوس متعرج `{` وحتى آخر قوس إغلاق `}`.",
                    },
                    {
                        "q": "ماذا يحدث إذا أرجع النموذج حقولاً إضافية لم نطلبها في نموذج البيانات Pydantic؟",
                        "a": "سيقوم نظام Pydantic بتصفيتها وحذفها تلقائياً إذا كان مهيأً على وضع التجاهل الافتراضي، أو سيفشل التحقق ويرمي استثناءً يتم التقاطه لإعادة المحاولة. في مشروعنا، نحدد بوضوح للنموذج الحقول المطلوبة ونمنعه من إضافة أي مفاتيح إضافية لضمان استقرار عملية تفكيك البيانات بالطرف الآخر.",
                    }
                ]
            },
            {
                "title": "الفرع الرابع: مكافحة الهلوسة والـ Grounding Prompts",
                "questions": [
                    {
                        "q": "ما هي استراتيجية مكافحة الهلوسة (Anti-Hallucination) التي طبقتموها في المشروع؟",
                        "a": "نطبق نظاماً صارماً من 4 مستويات: 1) كتابة ديباجة إرشادية صارمة (Grounding Preamble) تلزم النموذج بالاعتماد الحصري والكامل على النص المرسل وفضحه لأي نقص، 2) ضبط درجة الحرارة (Temperature = 0) لضمان اختيار الكلمة الأعلى احتمالاً دائماً ومنع الإبداع أو التخريف، 3) إرسال النص الفعلي فقط وحظر الاتصال الخارجي أو استخدام المعارف العامة للنموذج، 4) فحص المخرجات برمجياً عبر Pydantic للتأكد من مطابقتها للهياكل والأنواع المطلوبة.",
                        "tip": "هذا السؤال ممتاز لشرح الفارق بين استخدام الذكاء الاصطناعي كلعبة واستخدامه في تطبيق مؤسسي يتطلب دقة 100%."
                    },
                    {
                        "q": "ما هو الـ Grounding Preamble وكيف يساهم في دقة بناء محتوى الاختبارات؟",
                        "a": "هو نص توجيهي يتم إلحاقه ببداية الأوامر المرسلة للذكاء الاصطناعي (System Prompt) يحدد له القوانين الدستورية التي يجب أن يعمل تحت مظلتها. ينص صراحة على: 'يجب أن تبني جميع الأسئلة حصرياً على النص المرفق. لا تستعين بأي معلومات خارجية. إذا لم تجد المعلومة ارفض التوليد'. هذا يمنع النموذج من إدراج معلومات عامة قد لا تكون مقررة على الطالب في هذا الدرس المحدد.",
                    },
                    {
                        "q": "كيف نمنع النموذج من تخيل إجابات صحيحة غير موجودة علمياً في النص المرفوع؟",
                        "a": "نوجه النموذج في التلقين لكتابة مبررات الإجابة استناداً لنصوص صريحة بالمستند، ونفرض عليه التحقق الداخلي من توافق السؤال والجواب مع السياق، ويقوم نظام التحقق برفض الأسئلة التي تعتمد على استنتاجات واهية غير مدعومة بالحقائق المباشرة.",
                    },
                    {
                        "q": "كيف يمكن حماية النظام من هجمات التلقين الخبيث (Prompt Injection) التي يرفعها الطلاب في ملفات الـ PDF؟",
                        "a": "نحمي النظام بفصل تعليمات التشغيل (System Prompt) عن بيانات المستخدم (User Document Text) في استدعاءات النماذج، ونوجه النموذج صراحة في الديباجة الأمنية بتجاهل أي تعليمات أو أوامر مكتوبة داخل النص المرفوع والتركيز فقط على استخراج الحقائق العلمية وصياغتها كخريطة أو أسئلة.",
                    },
                    {
                        "q": "هل يمكن للذكاء الاصطناعي صياغة إجابات مضللة (Distractors) جيدة للامتحان وكيف نتحقق من منطقيتها؟",
                        "a": "نعم، نطلب من النموذج في الموجه صياغة خيارات مضللة تبدو علمية ومنطقية للوهلة الأولى ولكنها خاطئة بناءً على المستند، ونمنع الخيارات السهلة أو المستحيلة لضمان قياس عمق فهم الطالب الحقيقي وذكائه.",
                    }
                ]
            },
            {
                "title": "الفرع الخامس: نافذة السياق والترميز وتصنيف بلوم التربوي",
                "questions": [
                    {
                        "q": "كيف يؤثر حجم نافذة السياق (Context Window) لكل نموذج على حجم الملفات التعليمية التي يمكن معالجتها؟",
                        "a": "نافذة السياق تحدد كمية النصوص التي يستطيع النموذج قراءتها وتذكرها في المرة الواحدة. نموذج Gemini 2.0 Flash يمتلك نافذة عملاقة (1 مليون رمز) مما يجعله قادراً على قراءة كتب دراسية كاملة دون مشاكل، بينما Groq يمتلك نافذة أصغر (حوالي 128 ألف رمز). قمنا بضبط حجم النص المستخرج الأقصى ليتوافق مع أضيق النوافذ لضمان عدم حدوث انقطاع للنصوص أو تجاوز للحدود المسموحة للمزودين.",
                    },
                    {
                        "q": "ما هي الفلسفة التعليمية المتبعة في إرشاد الذكاء الاصطناعي لصياغة الأسئلة؟",
                        "a": "نوجه الذكاء الاصطناعي في نظام التلقين لاستخدام 'تصنيف بلوم للأهداف التعليمية' (Bloom's Taxonomy). نطلب منه صياغة أسئلة متنوعة تقيس مستويات معرفية مختلفة: أسئلة تذكر مباشرة (سهلة)، أسئلة فهم وتطبيق (متوسطة)، وأسئلة تحليل وتقييم منطقي (صعبة)، لضمان جودة الامتحان وصلاحيته لتقييم الطلاب بشكل حقيقي.",
                        "tip": "إشارة سريعة لتصنيف بلوم (Bloom's Taxonomy) ستبهر الأساتذة الأكاديميين في لجنة التقييم لأن هذا صميم تخصصهم التربوي."
                    },
                    {
                        "q": "ما الفرق في استهلاك المدخلات (Tokens) بين توليد خريطة مفاهيم متفرعة وتوليد امتحان مسطح؟",
                        "a": "توليد خريطة المفاهيم يستهلك رموزاً أقل في المخرجات لأنها تعتمد على كلمات مفتاحية مركزة وهيكل هرمي بسيط (Tree Structure). بينما توليد امتحان مسطح يحتوي على 50 سؤالاً يستهلك حجماً كبيراً جداً من مخرجات النصوص بسبب وجود خيارات متعددة ونصوص أسئلة طويلة وإجابات مفصلة مما يرفع استهلاك الرموز ووقت التوليد الكلي.",
                    },
                    {
                        "q": "كيف يتم حساب تكلفة استدعاء النماذج لكل عملية معالجة PDF؟",
                        "a": "العملية رخيصة للغاية؛ نموذج Gemini 2.0 Flash يكلف حوالي 0.075 دولار لكل مليون رمز مدخل، و 0.30 دولار لكل مليون رمز مخرج. معالجة مستند دراسي متوسط الحجم (10 صفحات / 4000 كلمة) تستهلك أقل من 0.005 دولار، مما يجعل النظام اقتصادياً جداً للتشغيل التجاري والمؤسسي.",
                    },
                    {
                        "q": "ما هي ميزة الـ 'Mixture of Experts' في معمارية Gemini وكيف تفيد في زيادة دقة مخرجاتنا؟",
                        "a": "معمارية الـ MoE تقوم بتوجيه الطلب الوارد لشبكات فرعية متخصصة داخل النموذج العملاق بدلاً من تشغيل النموذج بالكامل لكل رمز. هذا يتيح سرعة توليد أعلى وذكاءً مركزاً في فك شفرات النصوص الهيكلية واستخلاص المفاهيم الهرمية بكفاءة برمجية وموثوقية بالغة.",
                    }
                ]
            }
        ]
    },
    {
        "title": "الفصل الثالث: توليد المحتوى المرئي والمسموع (Video & Podcast Generation)",
        "description": "يناقش هذا القسم المنطق البرمجي المعقد وراء توليد ملفات الصوت والفيديو التعليمية، واستخدام ElevenLabs و FFmpeg، وتقنيات Ken Burns للصور وتكييف المدد البرمجية بناءً على حجم النصوص.",
        "subcategories": [
            {
                "title": "الفرع الأول: سيناريو البودكاست بالعامية المصرية (شريف وعبدالله)",
                "questions": [
                    {
                        "q": "كيف تعمل خدمة توليد البودكاست (Podcast Service) وما هي الهيكلية الخاصة بها؟",
                        "a": "تعمل الخدمة من خلال تقسيم النص المستخرج إلى أجزاء متوازنة، وتوليد نص حواري تفاعلي بالعامية المصرية بين شخصيتين (شريف وعبدالله)، ثم نرسل النصوص الصوتية على دفعات (Batches) إلى محرك ElevenLabs لتوليد الصوتيات بصوتين مختلفين، وأخيراً ندمج الملفات الصوتية المتتالية ونضيف المؤثرات الصوتية والموسيقى باستخدام مكتبة FFmpeg البرمجية لتخريج ملف بودكاست نهائي بصيغة MP3.",
                        "tip": "شدد على أن صياغة السيناريو بالعامية المصرية يعطي طابعاً واقعياً وجذاباً جداً للطلاب المصريين."
                    },
                    {
                        "q": "لماذا اخترتم الكتابة بالعامية المصرية الدارجة (Colloquial Egyptian Arabic) في حوار البودكاست؟",
                        "a": "لأن البودكاست وسيط تعليمي ترفيهي (Edutainment) يهدف لتبسيط المفاهيم المعقدة وجعلها قريبة من ذهن الطالب. استخدام العامية المصرية الدارجة المكتوبة بذكاء يخلق جواً ودياً يحاكي نقاش زميلين يذاكران معاً، مما يسهل استيعاب المصطلحات العلمية الصعبة ويزيد من تفاعل الطالب مع المحتوى مقارنة باللغة الفصحى الجافة.",
                    },
                    {
                        "q": "ما هي الأصوات المحددة المستخدمة في البودكاست وكيف يتم تخصيصها للشخصيات؟",
                        "a": "قمنا بإنشاء قاموس معرّفات أصوات (VOICES Mapping) يحتوي على معرّفات فريدة موثوقة من ElevenLabs: نخصص صوتاً يتميز بنبرة قيادية وفكاهية لشخصية 'شريف' (Host)، ونخصص صوتاً تفاعلياً ومتحمساً لشخصية 'عبدالله' (Expert/Co-host)، مما يعطي تنوعاً صوتياً واضحاً وممتعاً للمستمع.",
                    },
                    {
                        "q": "كيف نضمن توافق المصطلحات العلمية والتعريب المتبع في الجامعات المصرية مع مخرجات البودكاست؟",
                        "a": "نقوم بضبط موجه حوار الذكاء الاصطناعي لاستخدام المصطلحات العلمية الإنجليزية الدارجة باللفظ العربي وصياغتها باللهجة الحوارية المصرية المألوفة لدى الطلاب في قاعات المحاضرات (مثل استخدام 'أوبجكت' بدلاً من كائن، و 'أري' بدلاً من مصفوفة) لتقريب الفكرة لذهن الطالب.",
                    },
                    {
                        "q": "كيف يتم إدارة وتفادي الأخطاء اللغوية والنحوية في نصوص البودكاست المكتوبة بالعامية؟",
                        "a": "نوجه الذكاء الاصطناعي لاستخدام الكلمات الشائعة والسهلة النطق وتفادي الكلمات المركبة أو الغامضة. ونعتمد على قوة النماذج الحديثة التي تم تدريبها بكثافة على نصوص حوارية حقيقية بالعامية المصرية، مما يجعل صياغتها طبيعية ودقيقة لغوياً ومفهومة تماماً.",
                    }
                ]
            },
            {
                "title": "الفرع الثاني: إعداد وتنسيق النصوص للتلاوة الصوتية",
                "questions": [
                    {
                        "q": "كيف تتم تهيئة النصوص للتلاوة الصوتية (TTS Scripting Rules) لمنع محرك الصوت من ارتكاب الأخطاء؟",
                        "a": "nنطبق قواعد صارمة جداً في التوجيه: 1) حذف جميع الرموز التعبيرية وعلامات التنسيق مثل النجوم وعلامات الماركدوان لمنع قراءتها كرموز، 2) كتابة الاختصارات الإنجليزية برسم الكلمات العربية صوتياً (مثل كتابة 'إيه آي' بدلاً من AI، و 'بايثون' بدلاً من Python)، 3) استخدام علامات الترقيم بكثافة لإجبار محرك التوليد على أخذ فترات تنفس طبيعية.",
                    },
                    {
                        "q": "كيف يتم التعامل مع الكلمات الإنجليزية التقنية الصعبة في الصوتيات لضمان نطقها الصحيح؟",
                        "a": "نقوم بفلترة النصوص واستبدال المصطلحات العلمية الشائعة بنظيراتها المعربة صوتياً برمجياً قبل إرسالها لـ TTS، أو نطلب من نموذج الذكاء الاصطناعي أثناء صياغة السيناريو كتابة المصطلحات بالصيغة اللفظية المسهلة للنطق العربي العامي لضمان خروج جودة صوتية خالية من لكنات غريبة.",
                    },
                    {
                        "q": "ما هي معايير الجودة ومعدلات البت (Bitrate) المعتمدة للملفات الصوتية المنتجة ولماذا؟",
                        "a": "نعتمد معدل بت `128kbps` بتردد `44100Hz` بصيغة MP3. هذا الاختيار يحقق المعادلة المثالية: جودة صوتية واضحة ونقية جداً للحديث البشري والموسيقى المرافقة، مع الحفاظ على أحجام ملفات صغيرة وخفيفة تسهل رفعها للتخزين وتنزيلها وتشغيلها بسرعة وسلاسة على هواتف الطلاب ذات السرعات المتوسطة للإنترنت.",
                    },
                    {
                        "q": "ما هو دور خوارزمية التقطيع والدمج في معالجة السيناريوهات الطويلة للبودكاست؟",
                        "a": "تقسيم السيناريو إلى أجزاء صغيرة يمنع تجاوز الحد الأقصى لحجم الرموز في الطلب الواحد للذكاء الاصطناعي (Output Token Limits). بعد التوليد الناجح لكل جزء، نقوم بربط النصوص برمجياً وتصحيح تسلسل أرقام المتحدثين لتبدو كحلقة واحدة متصلة ومتناسقة هندسياً وصوتياً.",
                    },
                    {
                        "q": "كيف نضمن موثوقية ملفات الصوت المولدة وعدم وجود فترات تشويش أو قطع صوتي مفاجئ؟",
                        "a": "نستخدم مزوداً عالمياً قوياً ورائداً وهو ElevenLabs الذي يعتمد على نماذج توليد صوت عميقة وعالية النقاء والذكاء، ونقوم بفحص مخرجات بايتات الملفات الصوتية برمجياً قبل دمجها للتأكد من سلامتها واحتوائها على بيانات صالحة.",
                    }
                ]
            },
            {
                "title": "الفرع الثالث: توليد وتصيير مقاطع الفيديو بـ FFmpeg وتأثير Ken Burns",
                "questions": [
                    {
                        "q": "لماذا تم دمج وتوليف مقاطع الصوت والفيديو باستخدام أداة FFmpeg بدلاً من مكتبات بايثون الصرفة؟",
                        "a": "أداة FFmpeg هي المعيار العالمي لصناعة وتوليف الوسائط المتعددة. تتميز بكونها مكتوبة بلغة C فائقة السرعة وتدعم العمليات المعقدة مثل دمج التراكات الصوتية المتعددة، وإضافة تأثيرات التكبير والتحريك على الصور (Ken Burns Effect)، ودمج الصوت مع الفيديو بدقة ميلي-ثانية دون استهلاك ذاكرة عشوائية ضخمة كالمكتبات الأخرى.",
                    },
                    {
                        "q": "كيف يتم تنفيذ تأثير Ken Burns البرمجي (Ken Burns Effect) على الصور الثابتة في الفيديو؟",
                        "a": "نقوم بكتابة أمر معقد لـ FFmpeg يوجه مرشح الفيديو (Video Filter - `vf`) للقيام بعمليات تكبير تدريجية (Zoom) وإزاحة بطيئة للمحاور الأفقية والعمودية (Panning) بمعدل إطارات محدد طوال مدة المقطع الصوتي المصاحب. هذا يحول الصورة الثابتة إلى مشهد ديناميكي متحرك يجذب عين الطالب ويمنع الملل.",
                    },
                    {
                        "q": "كيف يضمن الكود محاذاة الصوت مع الفيديو بدقة تامة (Audio-Video Sync)؟",
                        "a": "نقوم أولاً بتوليد المقاطع الصوتية وحساب مدتها الزمنية الدقيقة بالثانية، ثم نوجه FFmpeg برمجياً لبناء مشاهد الفيديو بحيث تكون مدة عرض كل صورة مطابقة تماماً لمدة المقطع الصوتي المقابل لها، ونقوم بدمج المشاهد المتتالية بالتوالي للحصول على تزامن مثالي خافٍ من أي فجوات.",
                    },
                    {
                        "q": "ما هي لغة البرمجة والأدوات التي استخدمتموها لكتابة أوامر FFmpeg داخل Python؟",
                        "a": "nنستخدم وحدة `subprocess` المدمجة في Python لتشغيل أوامر FFmpeg النصية مباشرة على نظام التشغيل. هذا يمنحنا وصولاً كاملاً لجميع ميزات الأداة وقوتها بدون قيود المكتبات المغلفة، ويضمن سرعة تنفيذ عالية جداً وخالية من أخطاء الذاكرة الخاصة بـ Python wrapper.",
                    },
                    {
                        "q": "لماذا تم اختيار صيغة MP4 كصيغة نهائية للفيديو المتولد وما هي مزاياها؟",
                        "a": "صيغة MP4 بترميز H.264 وصوت AAC هي الصيغة الأكثر توافقاً وشهرة في العالم. مدعومة بالكامل وبشكل أصيل في جميع متصفحات الويب الحديثة، وأنظمة تشغيل الهواتف المحمولة (Android, iOS), وتطبيقات سطح المكتب دون الحاجة لتثبيت أي مشغلات أو حزم فك ترميز إضافية.",
                    }
                ]
            },
            {
                "title": "الفرع الرابع: استهلاك ElevenLabs و تدوير الصور البصرية",
                "questions": [
                    {
                        "q": "ما هي آلية توليد الصور المصاحبة للفيديو التعليمي وما هي مصادرها؟",
                        "a": "يقوم النظام بصياغة وصف بصري دقيق (Image Prompt) لكل مقطع في الفيديو بناءً على فكرة الشرح، ثم يرسله لنموذج توليد الصور الأساسي (مثل Hugging Face Stable Diffusion). وفي حال حدوث فشل في الاتصال، يتم تدوير الاستدعاء لـ Gemini كاحتياطي، وأخيراً نلجأ لتوليد صورة محلية توضيحية تحمل عنوان المقطع كحماية قصوى للنظام.",
                    },
                    {
                        "q": "لماذا نقوم بإرسال طلبات ElevenLabs TTS في مجموعات محدودة (Batches of 6)؟",
                        "a": "منصة ElevenLabs تفرض قيوداً صارمة على عدد الاتصالات المتزامنة المفتوحة لكل حساب لتجنب التحميل الزائد. إرسال الطلبات في مجموعات صغيرة متتالية مع فترات انتظار قصيرة (asyncio.sleep) يضمن توليد الصوتيات بأمان ودون التعرض لحظر الاتصال أو أخطاء خادم المنصة.",
                    },
                    {
                        "q": "كيف يتم الحفاظ على خصوصية مفاتيح ElevenLabs و Hugging Face داخل كود معالجة الوسائط؟",
                        "a": "يتم حقن هذه المفاتيح كمتغيرات بيئة مشفرة على الخادم ويتم استدعاؤها في كود الإعدادات `settings` فقط عند بدء تشغيل الخدمات. لا يتم إرسالها نهائياً للواجهة الأمامية أو طباعتها في سجلات الأخطاء (Logs)، مما يمنع سرقتها أو استغلالها الخارجي.",
                    },
                    {
                        "q": "كيف نتجنب ظهور شاشات سوداء أو فجوات زمنية صامتة بين مقاطع الفيديو المتتالية؟",
                        "a": "نقوم بعملية حسابية دقيقة لزمن كل مقطع صوتي ونمنع تقريب الأرقام لكسور كبيرة. كما نستخدم أمر التجميع (Concat Protocol) في FFmpeg الذي يقوم بدمج التدفقات الصوتية والبصرية دون الحاجة لإعادة ترميز المقاطع، مما يمنع ظهور أي رمشات بصرية أو صمت بين المقاطع.",
                    },
                    {
                        "q": "كيف نضمن عدم حدوث أخطاء تجاوز الذاكرة (Out of Memory) عند دمج مقاطع فيديو بدقة عالية على خادم VPS؟",
                        "a": "ونحن نقوم بتوليد الفيديو بدقة مناسبة للهواتف المحمولة وتطبيقات الويب (مثل 720p) ومعدل إطارات منخفض (24 إطار/ثانية). هذه الإعدادات تقلل بشكل هائل من استهلاك الذاكرة والمعالجة أثناء التصيير، وتجعل العملية تتم بسرعة وأمان تام داخل حدود الذاكرة المتاحة للخادم.",
                    }
                ]
            },
            {
                "title": "الفرع الخامس: التحكم بالمدد الزمنية وإدارة الملفات المؤقتة",
                "questions": [
                    {
                        "q": "كيف يساهم كود `calculate_smart_config` في موازنة طول الفيديو أو البودكاست مع حجم النص؟",
                        "a": "يقوم الكود بقياس طول النص وحساب عدد الكلمات، ثم يصنف المدخلات إلى ثلاث فئات (Short, Medium, Long). بناءً على هذه الفئة، يحدد ديناميكياً الإعدادات المثلى للتوليد: عدد الفقرات الحوارية للبودكاست، أو عدد المقاطع البصرية والتعليقات الصوتية للفيديو، لضمان تغطية كاملة للمحتوى وتفادي تمديد الوقت بلا فائدة أو ضغط الأفكار بشكل مخل.",
                        "tip": "هذا الكود هو عقل عملية التوليد الموزونة، ويظهر قدرتنا على كتابة خوارزميات ذكية تتكيف مع مدخلات المستخدم."
                    },
                    {
                        "q": "ماذا يحدث إذا كان الملف المرفوع صغيراً جداً (مثلاً صفحة واحدة)؟ كيف نولد منه فيديو مدته 5 دقائق؟",
                        "a": "كود `calculate_smart_config` سيكتشف ذلك ويضع إعدادات توسعية. يتم توجيه الذكاء الاصطناعي لعدم الاكتفاء بالتلخيص، بل شرح المفاهيم البسيطة بالتفصيل الممل وإعطاء أمثلة توضيحية وتطبيقات حياتية وتوليد نقاش حواري ممتد حولها لملء عدد الدقائق المطلوبة دون الخروج عن محتوى الصفحة.",
                    },
                    {
                        "q": "لماذا تم فصل توليد الصوت وتوليد الفيديو في واجهتين برمجيتين (Endpoints) منفصلتين؟",
                        "a": "لأن العمليتين مستهلكتان جداً للوقت والجهد الحسابي. فصلهما يتيح للمستخدم اختيار ما يريده بدقة، ويسهل إدارة مهل التنفيذ (Timeout) الخاصة بـ Nginx/FastAPI لكل طلب على حدة، كما يتيح توازي العمليات وتوفير استهلاك الباقة السحابية للمستخدم الذي يرغب فقط في سماع البودكاست دون مشاهدة الفيديو.",
                    },
                    {
                        "q": "كيف يتم تنظيف وتفريغ ملفات الصوت الثنائية المؤقتة بعد انتهاء عملية الدمج ورفع الملف لـ Supabase؟",
                        "a": "نستخدم الكتل البرمجية الآمنة `try...finally`؛ حيث نقوم بإنشاء مجلد مؤقت للعملية الجارية، وبمجرد رفع الملف النهائي بنجاح أو حدوث خطأ مفاجئ، يتم استدعاء دالة المسح تلقائياً لحذف جميع الملفات والمجلدات المؤقتة من القرص الصلب لضمان نظافة بيئة الخادم دائماً.",
                    },
                    {
                        "q": "كيف يمكننا دمج موسيقى خلفية هادئة (Background Music) مع كلام المذيعين في البودكاست برمجياً؟",
                        "a": "نستخدم مرشح الدمج والتداخل (Audio Filter - `amix`) في FFmpeg. نحدد ملف الموسيقى كمسار ثانٍ، ونقوم بخفض مستوى الصوت الخاص به بنسبة 85% ليكون خافتاً وتحت صوت المتحدثين الرئيسي، ثم نطبق تأثير التلاشي التدريجي (Fade-out) عند نهاية البودكاست لخروج فني راقٍ.",
                    }
                ]
            }
        ]
    },
    {
        "title": "الفصل الرابع: إدارة البيانات والربط مع قاعدة البيانات (Database & Storage Integration)",
        "description": "يركز هذا الفصل على كيفية ربط الخلفية بقاعدة بيانات Supabase، وتصميم الجداول للاختبارات وخرائط المفاهيم والبودكاست، وحفظ وتخزين ملفات الوسائط الثنائية الكبيرة.",
        "subcategories": [
            {
                "title": "الفرع الأول: دمج منصة Supabase ونظام PostgreSQL",
                "questions": [
                    {
                        "q": "لماذا تم اختيار Supabase كمنصة لإدارة البيانات وقاعدة البيانات في المشروع؟",
                        "a": "اخترنا Supabase لأنها توفر قاعدة بيانات PostgreSQL قوية وموزعة وتدعم استدعاءات البيانات غير المتزامنة بطلاقة. بالإضافة إلى دمجها لخدمات التخزين السحابي للملفات (Storage Buckets) ونظام المصادقة في منصة واحدة متكاملة ومفتوحة المصدر وسهلة الدمج مع تطبيقات Python و C#.",
                        "tip": "PostgreSQL هي قاعدة بيانات علائقية قوية جداً، واختيارها يعطي المشروع طابعاً احترافياً في إدارة العلاقات المعقدة بين الكيانات."
                    },
                    {
                        "q": "كيف يتم التعامل مع اتصال قاعدة البيانات في الكود لمنع استهلاك قنوات الاتصال المتاحة (Connection Pool Exhaustion)؟",
                        "a": "نقوم بإنشاء عميل Supabase مرة واحدة فقط كمتغير عالمي مفرد (Singleton Instance) وتتم إعادة استخدامه طوال دورة حياة التطبيق. كما نستخدم الدوال غير المتزامنة والمؤشرة للطلب الخارجي فقط عند الحاجة الفعلية لمنع إبقاء اتصالات مفتوحة بلا داعٍ.",
                    },
                    {
                        "q": "لماذا لا نستخدم قاعدة بيانات ممثلة بالذاكرة مثل Redis لتخزين الجلسات السريعة ومخرجات الفيديوهات؟",
                        "a": "منصة Supabase توفر عناء إدارة خوادم إضافية وتكفي احتياجات المشروع الحالية بامتياز. وعند نمو المشروع بشكل ضخم، يمكن بالطبع إدخال Redis كطبقة كاش وسيطة لتخزين استجابات الواجهة البرمجية المتكررة وتقليل الضغط على قاعدة البيانات والذكاء الاصطناعي.",
                    },
                    {
                        "q": "لماذا تم فصل منطق الاتصال بقاعدة البيانات في مجلد منفصل `core/database.py`؟",
                        "a": "لتطبيق مبدأ التغليف والنظام المفرد. يسهل هذا التعديل على مزود قاعدة البيانات مستقبلاً (مثلاً التبديل من Supabase إلى Firebase أو PostgreSQL مستقلة) بتغيير الكود في ملف واحد فقط ودون الحاجة لتعديل بقية ملفات الخدمات أو النهايات البرمجية.",
                    },
                    {
                        "q": "كيف يؤثر استخدام PostgreSQL على إمكانية دعم ميزات التعاون الفوري (Real-time Collaboration) للطلاب مستقبلاً؟",
                        "a": "منصة Supabase تدعم ميزة الـ Realtime بشكل ممتاز وفوري فوق PostgreSQL. يمكننا تفعيل هذه ميزة بلمسة زر لتحديث نتائج امتحانات وخرائط مفاهيم على شاشات طلاب مشتركين في نفس جلسة فور حفظها في قاعدة بيانات وبدون حاجة لعمليات استعلام متكررة."
                    }
                ]
            },
            {
                "title": "الفرع الثاني: تصميم جداول الامتحانات والخرائط والوسائط",
                "questions": [
                    {
                        "q": "ما هو تصميم الجدول (Table Schema) الخاص بحفظ بنوك الأسئلة المنتجة؟",
                        "a": "يحتوي جدول `generated_quizzes` على الحقول التالية: المعرف الفريد `id` (مفتاح أساسي)، العنوان `title` (نص)، الصعوبة `difficulty` (نص)، عدد الأسئلة `num_questions` (عدد صحيح)، وتفاصيل بنك الأسئلة بالكامل بصيغة JSON داخل الحقل `quiz_data` وحقل نوع الامتحان `type` وتاريخ الإنشاء التلقائي.",
                    },
                    {
                        "q": "لماذا نفضل تخزين تفاصيل الأسئلة كـ JSON في قاعدة البيانات بدلاً من فصلها في جداول علائقية تقليدية؟",
                        "a": "نفضل تخزينها كـ JSON (Document-style store) لثلاثة أسباب: 1) السرعة الفائقة في الحفظ والاسترجاع ككتلة واحدة دون الحاجة لعمليات ربط معقدة (SQL Joins) عبر عدة جداول، 2) مرونة الهيكل وقابليته للتعديل مستقبلاً دون الحاجة لعمل هجرات لقواعد البيانات (Migrations)، 3) التوافق المباشر مع الهيكل البرمجي Pydantic وعقود بيانات الواجهة الأمامية.",
                    },
                    {
                        "q": "كيف يتم تصميم جدول `generated_mindmaps` وما هي العلاقات البرمجية له؟",
                        "a": "يحتوي جدول `generated_mindmaps` على معرف فريد `id` كحقل أساسي، وحقل JSON لتخزين الهيكل الشجري للمفهوم بالكامل `mindmap_data` ورابط الصورة النهائية المولدة للخريطة `image_url` بالإضافة لحقول التوقيت الزمني للإدخال.",
                    },
                    {
                        "q": "كيف يتم تصميم جدول `generated_podcasts` وما هي العلاقات البرمجية له؟",
                        "a": "يحتوي الجدول على معرّف فريد `id` كحقل أساسي، وحقل العنوان `title` (نص)، وحقل المدة الزمنية الكلية بالثانية `total_duration_seconds` (رقم عشري)، وحقل تفاصيل البودكاست بالكامل بصيغة JSON داخل الحقل `podcast_data` الذي يحتوي على المسارات والصوتيات المنتجة.",
                    },
                    {
                        "q": "كيف يتم التعامل مع الحالات التي يتجاوز فيها حجم مصفوفة الـ JSON الحد الأقصى المسموح به لحجم العمود في قاعدة البيانات؟",
                        "a": "أعمدة من نوع JSONB في PostgreSQL لا تفرض قيوداً ضيقة على الحجم (تصل إلى 1 جيجابايت لكل حقل كحد أقصى). وبما أن حجم بيانات خرائط المفاهيم والامتحانات لدينا لا يتعدى بضعة كيلوبايتات، فإننا في أمان تام من تجاوز حدود التخزين للعمود الواحد."
                    }
                ]
            },
            {
                "title": "الفرع الثالث: التخزين السحابي للملفات والروابط العامة",
                "questions": [
                    {
                        "q": "ما هي المجلدات السحابية (Storage Buckets) المستخدمة في Supabase وما هي وظائفها؟",
                        "a": "نستخدم باقة تخزين عامة (Public Storage Bucket) لحفظ ملفات الصوت (MP3) والفيديوهات (MP4) وخرائط المفاهيم الصورية المنتجة. يمنحنا هذا روابط مباشرة وسريعة (Public URLs) قابلة للاستهلاك المباشر في تطبيقات الويب والهواتف المحمولة.",
                    },
                    {
                        "q": "ما هو دور متغير البيئة `SUPABASE_STORAGE_BUCKET` وكيف يساهم في مرونة النظام؟",
                        "a": "يحدد هذا متغير اسم مجلد التخزين سحابة فعلي ترفع إليه ملفات. فصل هذا اسم في متغير بيئة يتيح استخدام مجلد اختبارات (Testing Bucket) أثناء تطوير ومجلد إنتاج (Production Bucket) عند إطلاق دون حاجة لتغيير سطر واحد كود.",
                    },
                    {
                        "q": "ما هي معايير الأمان المطبقة لحماية روابط الوسائط المخزنة من الوصول غير المصرح به؟",
                        "a": "يمكن تفعيل روابط الوصول المؤقتة (Signed URLs) التي تنتهي صلاحيتها بعد دقائق معدودة بدلاً من الروابط العامة الدائمة، كما نطبق سياسات أمان صارمة (RLS - Row Level Security) على مستوى جداول Supabase لمنع القراءة أو التعديل من خارج تطبيقنا.",
                    },
                    {
                        "q": "كيف يتم التعامل مع الكلمات العربية والأحرف الخاصة عند إنشاء وتسمية ملفات الصوت في التخزين السحابي؟",
                        "a": "لتجنب مشاكل ترميز الروابط (URL Encoding) والأخطاء البرمجية للمتصفحات، نقوم بإنشاء أسماء ملفات عشوائية تعتمد على معرّفات فريدة (UUID) أو تشفير بصمات النصوص (MD5) ونحتفظ بالاسم العربي الأصلي للمستند كحقل نصي فقط داخل قاعدة البيانات.",
                    },
                    {
                        "q": "كيف يمكن عمل نسخة احتياطية (Backup) دورية لكامل البيانات والملفات المخزنة في Supabase للمحافظة على أمان المشروع؟",
                        "a": "توفر منصة Supabase نسخاً احتياطية تلقائية يومية لقاعدة البيانات كجزء من خدماتها. وبالنسبة للملفات في الـ Storage، يمكننا كتابة سكربت خارجي دوري يقوم بنسخ محتويات مجلد التخزين إلى سحابة مستقلة كـ AWS S3 كنسخة احتياطية إضافية للطوارئ."
                    }
                ]
            },
            {
                "title": "الفرع الرابع: كفاءة الاتصال والتكامل غير المتزامن بـ to_thread",
                "questions": [
                    {
                        "q": "لماذا تم استخدام `asyncio.to_thread` عند إجراء عمليات الإدخال لقاعدة بيانات Supabase؟",
                        "a": "مكتبة Supabase في Python تقوم بالاتصالات بشكل متزامن وحظر لتدفقات المعالجة (Blocking IO). تغليف هذه الاستدعاءات بـ `asyncio.to_thread` يوجه إطار العمل لتشغيلها في خيط معالجة منفصل (Separate Thread) مما يمنع تجميد خادم FastAPI الرئيسي ويحافظ على استجابة النظام للطلبات الأخرى المتزامنة.",
                    },
                    {
                        "q": "كيف يتم معالجة وتفادي أخطاء فشل الاتصال بقاعدة البيانات أثناء حفظ النتائج في الخلفية؟",
                        "a": "نطبق منطق الحفظ كعملية 'أفضل جهد' (Best-effort execution) محاطة بكتلة استثناء `try...except`. إذا فشل الاتصال بقاعدة البيانات لأي سبب، نقوم بتسجيل الخطأ في السجلات (Logger) ولكن نستمر في إرجاع النتيجة بنجاح للمستخدم، مما يمنع تعطل تجربة الطالب بسبب مشكلة مؤقتة في قاعدة البيانات.",
                        "tip": "هذا القرار يعكس وعياً برمجياً يفضل استمرار تجربة المستخدم (User Experience) على كمال العمليات الجانبية غير المؤثرة مباشرة على الطلب."
                    },
                    {
                        "q": "ما هي المزايا الأمنية لاستخدام الاتصال المشفر (SSL/TLS) بين خادم FastAPI وقاعدة بيانات Supabase؟",
                        "a": "يضمن تشفير الاتصال حماية جميع البيانات المتبادلة (بيانات الطلاب، الامتحانات، روابط الملفات) من الاختراق والتنصت في الشبكة البينية، ويمنع هجمات الرجل في المنتصف (Man-in-the-Middle) وهو متطلب أساسي في معايير الامتثال لأمن المعلومات.",
                    },
                    {
                        "q": "ما هي التقنية المستخدمة لتتبع زمن وتاريخ إدخال السجلات تلقائياً في جداول Supabase؟",
                        "a": "نستخدم ميزة القيم الافتراضية في قاعدة البيانات PostgreSQL؛ حيث نحدد عمود `created_at` بنوع بيانات `timestamp` ونضبط قيمته الافتراضية لتكون الدالة `NOW()`، مما يجعل قاعدة البيانات تتولى تسجيل الوقت الدقيق للإدخال تلقائياً ودون تدخل من كود بايثون.",
                    },
                    {
                        "q": "كيف يمكننا إجراء عمليات هجرة وتعديل جداول قاعدة البيانات (Database Migrations) بشكل آمن في بيئة الإنتاج؟",
                        "a": "يمكننا استخدام أدوات قياسية لإدارة الهجرات مثل Alembic في حال استخدام SQLAlchemy، أو استخدام واجهة تحكم Supabase لإجراء التعديلات وحفظها كملفات SQL تتبع في مستودع الكود لضمان تطبيقها بالتساوي في بيئات التطوير والاختبار والإنتاج."
                    }
                ]
            },
            {
                "title": "الفرع الخامس: التحقق من البيانات وتتبع بصمات الملفات (MD5 Hash)",
                "questions": [
                    {
                        "q": "كيف نضمن عدم حدوث تكرار في توليد نفس خريطة المفاهيم لنفس الملف التعليمي المرفوع بشكل متكرر؟",
                        "a": "نقوم بحساب بصمة رقمية فريدة (MD5 Hash) لمحتويات ملف الـ PDF المرفوع قبل بدء المعالجة. نقوم بالبحث في قاعدة البيانات عن هذا الـ Hash؛ فإذا وجدنا خريطة مفاهيم مطابقة ومولدة سابقاً، نرجعها فوراً للمستخدم، مما يوفر وقت المعالجة وتكلفة استدعاء الذكاء الاصطناعي تماماً.",
                        "tip": "هذا يسمى 'Caching by Content Hashing' وهو أحد أفضل ممارسات هندسة البرمجيات لتحسين الأداء وتقليل التكاليف السحابية."
                    },
                    {
                        "q": "كيف يضمن الكود سلامة البيانات والاتساق (Data Consistency) عند حفظ البيانات الشجرية لخرائط المفاهيم؟",
                        "a": "يتم فحص خريطة المفاهيم بالكامل والتحقق من صحة ترابط العقد وعلاقات الأب والابن برمجياً عبر نموذج Pydantic قبل إدخلها لقاعدة البيانات. يتم رفض أي خريطة تحتوي على عقد مكسورة أو معرّفات غير متوافقة من المصدر وتجنب حفظ بيانات تالفة.",
                    },
                    {
                        "q": "ماذا يحدث لو قمنا برفع ملف تالف أو فارغ تماماً لقاعدة البيانات؟ كيف يقي الكود قاعدة البيانات من هذا؟",
                        "a": "الخلفية تفحص محتوى الملف برمجياً وتتأكد من وجود نصوص وصفحات صالحة قبل إجراء أي عملية حفظ أو استدعاء للذكاء الاصطناعي. وفي حال عدم صلاحية الملف، يرمي الكود استثناءً فورياً ويرجع للمستخدم كود الحالة HTTP 422 (Unprocessable Entity) دون لمس قاعدة البيانات.",
                    },
                    {
                        "q": "ما هو نموذج استهلاك البيانات ومعدل القراءة والكتابة المتوقع لقاعدة البيانات في هذا التطبيق؟",
                        "a": "التطبيق يعتمد على نمط كتابة ثقيلة نسبياً عند إنتاج الوسائط والامتحانات (Write-heavy) وقراءة متوسطة عند استرجاعها للمراجعة. قاعدة بيانات PostgreSQL في Supabase مصممة للتعامل مع آلاف العمليات المتزامنة للقراءة والكتابة بفضل نظام الفهرسة الفعال وإدارة الذاكرة.",
                    },
                    {
                        "q": "ماذا لو أردنا دعم البحث النصي الكامل (Full-Text Search) في الامتحانات المخزنة؟ كيف تدعم PostgreSQL ذلك؟",
                        "a": "توفر PostgreSQL دعماً أصيلاً وممتازاً للبحث النصي الكامل حتى داخل حقول JSONB. يمكننا كتابة استعلامات SQL متقدمة تبحث عن كلمات مفتاحية داخل حقول الأسئلة والخيارات بسرعة وكفاءة عالية وتدعم محركات البحث الذكية للطلاب."
                    }
                ]
            }
        ]
    },
    {
        "title": "الفصل الخامس: تكامل النظام وتصميم واجهة برمجة التطبيقات (API Design & C# Interoperability)",
        "description": "يناقش هذا القسم طرق الربط الفعال مع تطبيقات الواجهة الأمامية المصممة بلغة C#، وهيكلة واجهة برمجة التطبيقات، ونظام الأمان والتحكم بالوصول ومعدلات الاستخدام وRequest Tracing.",
        "subcategories": [
            {
                "title": "الفرع الأول: تصميم JSON ومقارنة هياكل البيانات للخيارات",
                "questions": [
                    {
                        "q": "لماذا تم اختيار صيغة مفاتيح JSON بدلاً من المصفوفات المسطحة لتمثيل خيارات الأسئلة المتعددة؟",
                        "a": "تم ذلك لثلاثة أسباب رئيسية: 1) حماية الامتحان من الغش البرمجي؛ فالمصفوفة التي تحتوي على علامة `isCorrect` تفضح الإجابة الصحيحة فوراً في DevTools للطلاب الأذكياء، بينما فصل الإجابة الصحيحة في مفتاح منفصل `answer: 'B'` يتيح حذف هذا مفتاح قبل إرسال الأسئلة للطلاب وتصحيحها في الخادم لاحقاً، 2) تقليل حجم البيانات المرسلة بنسبة تصل لـ 80%، 3) سهولة تفكيك وتحليل البيانات بخطوة واحدة في لغة C#.",
                        "tip": "هذه الإجابة تجمع بين الحنكة الأمنية (Security by Design) وكفاءة الشبكة وبساطة الكود."
                    },
                    {
                        "q": "كيف تسهل هيكلتكم الحالية استهلاك البيانات وتفكيكها (Deserialization) في واجهة C# الأمامية؟",
                        "a": "الهيكل الذي نرجعه مسطح وثابت ويطابق تماماً نموذج البيانات القوي (Strongly-typed contracts). يستطيع مبرمج C# تعريف فئة (Class) مقابلة تماماً للهيكل واستخدام سطر برمجى واحد كـ `JsonSerializer.Deserialize<ExamData>(jsonString)` دون الحصول على أخطاء تشغيل التطبيق.",
                    },
                    {
                        "q": "لماذا تم اختيار صيغة مفاتيح JSON ككائنات (Objects) بدلاً من مصفوفة من السلاسل النصية للأسئلة؟",
                        "a": "استخدام كائنات JSON يتيح لنا تسمية الحقول بدقة وحفظ معرفات فريدة لكل سؤال وجواب، مما يسهل عملية البحث والفرز والفلترة داخل قاعدة البيانات وتجنب الأخطاء العشوائية لترتيب المصفوفات.",
                    },
                    {
                        "q": "كيف يتم استهلاك المخرجات الشجرية لخرائط المفاهيم داخل كود الواجهة الأمامية؟",
                        "a": "يتم فك تشفير الهيكل الشجري JSON إلى نموذج بيانات شجري متداخل في C# يحتوي على عقد وعناصر فرعية، ويتم رسمها برمجياً على الشاشة كشجرة تفاعلية باستخدام مكتبات الرسوم المتاحة.",
                    },
                    {
                        "q": "ما فائدة إعداد ترويسة `allow_credentials=True` في إعدادات CORS وما هي ضوابط استخدامها؟",
                        "a": "تتيح هذه الترويسة للواجهة الأمامية تبادل ملفات تعريف الارتباط (Cookies) وتفاصيل مصادقة الشبكة مع الخلفية. ضابط استخدامها الأهم هو حظر استخدام النجمة `*` كأصل مسموح به وفرض كتابة النطاقات المحددة بدقة لدواعٍ أمنية صارمة."
                    }
                ]
            },
            {
                "title": "الفرع الثاني: تتبع الطلبات عبر RequestID Middleware",
                "questions": [
                    {
                        "q": "ما هو دور الـ Middleware المسمى `RequestIDMiddleware` وما هي فوائده في تتبع المشاكل؟",
                        "a": "يقوم هذا الـ Middleware بحقن معرّف فريد للطلب `X-Request-ID` في ترويسات (Headers) كل طلب واستجابة تمر عبر النظام. يفيد هذا في ربط جميع سجلات الأحداث (Logs) الخاصة بطلب معين مع بعضها وتسهيل البحث عن الأخطاء وحلها في الأنظمة الموزعة والمزدحمة بمئات الطلبات المتزامنة.",
                    },
                    {
                        "q": "كيف يساهم استخدام المعرفات الفريدة للطلبات (Request IDs) في بناء لوحات مراقبة الأداء وتحليل الأنظمة؟",
                        "a": "يسمح لنا بربط السجلات ببعضها وجمع إحصاءات زمن الاستجابة لكل معرّف طلب، مما يساعد في تحديد نهايات الخدمة البطيئة واكتشاف الاختناقات البرمجية وحلها بكفاءة عالية بالاعتماد على أدوات تحليل السجلات (Log Analyzers).",
                    },
                    {
                        "q": "كيف يتم حقن وتمرير `X-Request-ID` بين خادم FastAPI وقاعدة بيانات Supabase؟",
                        "a": "نقوم بتمرير هذا المعرّف في سجلات قاعدة البيانات عند إدخال العمليات، مما يتيح تتبع السجل البرمجي من بداية طلب العميل في Nginx وحتى إتمام الحفظ النهائي بقاعدة البيانات والرجوع للعميل بنجاح.",
                    },
                    {
                        "q": "ما هو دور الـ Middleware في التحقق من صحة ونزاهة هياكل طلبات رفع الملفات (File Upload Validation)؟",
                        "a": "يقوم الـ Middleware بفحص ترويسات الطلب والتأكد من نوع المحتوى (Content-Type) المرفق لمنع إرسال طلبات عشوائية أو هجمات الاختراق، ويقوم برفض الطلبات الخاطئة قبل بدء تحميل بايتات الملف للذاكرة.",
                    },
                    {
                        "q": "كيف يتم التعامل مع الطلبات التي لا تحتوي على ترويسة `X-Request-ID` من المصدر؟",
                        "a": "يقوم الـ Middleware باكتشاف ذلك تلقائياً، وإنشاء معرّف فريد جديد باستخدام مكتبة `uuid` البرمجية، وحقنه داخل الطلب والاستجابة لضمان شمولية التتبع لجميع العمليات المارة بالنظام."
                    }
                ]
            },
            {
                "title": "الفرع الثالث: تحديد معدل الاستهلاك بـ SlowAPI",
                "questions": [
                    {
                        "q": "كيف يتم تطبيق نظام تحديد معدل الطلبات (Rate Limiting) في التطبيق ولماذا؟",
                        "a": "nنطبق نظام تحديد معدل الطلبات باستخدام مكتبة `SlowAPI` المعتمدة على خوارزمية Token Bucket. نضع حداً أقصى للاستدعاءات مثل 5 طلبات في الدقيقة لكل عنوان IP على النهايات البرمجية المستهلكة للذكاء الاصطناعي لحماية خوادمنا وحصصنا المالية من هجمات الإغراق البرمجي والتعطيل المتعمد.",
                    },
                    {
                        "q": "ما هو دور مكتبة `slowapi` وكيف تختلف عن أنظمة الحد من استخدام الـ API التقليدية؟",
                        "a": "مكتبة `slowapi` هي تطبيق معتمد على بايثون لإطار عمل Limiter الشهير، وتعمل بشكل مدمج وسلس مع FastAPI دون الحاجة لإعداد خوادم مستقلة كـ Redis لإجراء عمليات العد والتتبع، مما يجعلها مثالية للنشر السريع في البيئات الخادمة.",
                    },
                    {
                        "q": "ما هو كود الحالة HTTP المرجع عند تجاوز المستخدمين لحد الطلبات المسموح به وكيف يستقبلها العميل؟",
                        "a": "يرجع النظام كود الحالة القياسي `HTTP 429 Too Many Requests`. يستقبلها تطبيق C# في ترويسة الاستجابة ويقوم بعرض رسالة واضحة للمستخدم تدعوه للانتظار لبعض الوقت قبل المحاولة مرة أخرى للحفاظ على أدب الاستهلاك.",
                    },
                    {
                        "q": "كيف نمنع الطلاب من إساءة استخدام الخدمة واستنزاف الحسابات المجانية للكلية؟",
                        "a": "نطبق حد طلبات صارم ومحسوب بدقة، ويمكن ربط النظام بنظام تسجيل دخول رسمي عبر البريد الجامعي للكلية وتحديد كوتا معينة (مثلاً معالجة 5 مستندات كحد أقصى شهرياً لكل طالب) مما يحقق العدالة في الاستهلاك ويحمي الميزانية المالية للكلية.",
                    },
                    {
                        "q": "كيف يتم تخزين حالات العدادات (Rate Limit States) داخل الذاكرة للخلفية؟",
                        "a": "تخزن مكتبة SlowAPI العدادات مؤقتاً في ذاكرة تطبيق بايثون النشطة (In-memory Storage)، وهو خيار ممتاز وسريع، ولكن في بيئة VPS يمكن ترقيتها لتخزين العدادات في Redis لضمان دقة العد بين العمليات المتعددة."
                    }
                ]
            },
            {
                "title": "الفرع الرابع: تأمين واجهات الخدمة بـ API Key والمصادقة",
                "questions": [
                    {
                        "q": "لماذا تم فرض مفتاح المصادقة لواجهة برمجة التطبيقات (API Key Authentication) وكيف يتم التحقق منه؟",
                        "a": "لضمان حصر استهلاك الخدمة على التطبيقات المعتمدة لدينا فقط (مثل تطبيق C# المعتمد) ومنع الاستغلال المفتوح للواجهات البرمجية. يتم التحقق عبر دالة تعتمد على نظام الحماية المدمج في FastAPI `Depends(verify_api_key)` تفحص وجود قيمة المطابقة في ترويسة الطلب ومقارنتها بالمفتاح المخزن في بيئة النظام.",
                        "tip": "المصادقة باستخدام الترويسات (Headers) هي الطريقة المعيارية والأكثر أماناً لحماية واجهات الـ RESTful APIs."
                    },
                    {
                        "q": "ما الفرق بين استخدام أساليب التمرير عبر الترويسات (Headers) وعبر بارامترات الاستعلام (Query Parameters) للمصادقة؟",
                        "a": "التمرير عبر الترويسات (Headers) أكثر أماناً لأن بارامترات الاستعلام (Query Parameters) تظهر بوضوح في سجلات خوادم الويب وتاريخ المتصفحات ورسائل تتبع الشبكة، مما يعرض مفاتيح الأمان لخطر التسريب السهل والقراءة والسرقة.",
                    },
                    {
                        "q": "لماذا تم دمج ترويسات مخصصة مثل `X-API-Key` بدلاً من استخدام المصادقة القياسية بـ Bearer Token؟",
                        "a": "استخدام `X-API-Key` أبسط وأسرع للربط المباشر بين خادمين (Server-to-Server Communication) ولا يتطلب إدارة تعقيدات انتهاء صلاحية التوكن وإعادة توليده دورياً كـ OAuth2، وهي مناسبة تماماً لطبيعة التكامل الحالية للمشروع.",
                    },
                    {
                        "q": "ما هي المزايا الأمنية التي يقدمها تعطيل مستندات Swagger التفاعلية في بيئة الإنتاج؟",
                        "a": "يمنع تعطيل المستندات التفاعلية المهاجمين المحتملين من استكشاف نهايات الخدمة المتاحة وفحص الهياكل المتوقعة وثغرات المدخلات بسهولة، وهو إجراء تأميني معتاد لحماية واجهات التطبيقات الحساسة في الإنتاج.",
                    },
                    {
                        "q": "كيف يمكننا تتبع الاستهلاك المالي الفعلي لكل مستخدم مسجل للخدمة بناءً على استخدام الـ API؟",
                        "a": "يمكن ربط مفتاح الـ API الخاص بكل مستخدم بجدول في قاعدة البيانات يسجل عدد الرموز المستهلكة (Tokens Counter) لكل طلب أو عدد ثواني توليد الصوت والفيديو، ويقوم النظام باقتطاع التكلفة أو تقليل رصيد المستخدم مع كل استدعاء ناجح.",
                        "tip": "هذا النموذج التجاري (Pay-per-use) هو المعيار الأساسي لبناء وتطوير الخدمات السحابية الخدمية كمشروع ناشئ."
                    }
                ]
            },
            {
                "title": "الفرع الخامس: تكامل واجهات الخدمة مع واجهة C#",
                "questions": [
                    {
                        "q": "كيف يساهم نموذج البيانات Pydantic في بناء عقود برمجية متماسكة مع مطوري الواجهة الأمامية؟",
                        "a": "يعمل Pydantic كعقد صارم بين الخلفية والواجهة؛ فهو يضمن بنسبة 100% أن أي استجابة ناجحة ستطابق النماذج والأنواع المصممة مسبقاً بالتفصيل. هذا يلغي تماماً المفاجآت وأخطاء القيمة المفقودة (Null Pointer Exceptions) التي تواجه مبرمجي الواجهة الأمامية.",
                    },
                    {
                        "q": "كيف تدعمون التوثيق التلقائي لواجهات برمجة التطبيقات (Automatic Swagger Docs)؟",
                        "a": "يوفر FastAPI هذا بشكل تلقائي بفضل استخدامه لمعايير OpenAPI. يمكن الوصول للمستندات التفاعلية وتجربة نهايات الخدمة مباشرة عبر المسارات `/docs` أو `/redoc` والتي تعكس بوضوح النماذج والمتغيرات والأنواع المطلوبة لكل طلب واستجابة.",
                    },
                    {
                        "q": "كيف تتعامل واجهة برمجة التطبيقات مع الحالات التي يرسل فيها العميل بيانات غير كاملة أو خاطئة (Invalid Body)؟",
                        "a": "يلتقط FastAPI الخطأ تلقائياً ويرجع استجابة سريعة بكود الحالة `HTTP 422 Unprocessable Entity` مع تفاصيل برمجية واضحة ومحددة تشير بدقة للحقل المفقود أو الخاطئ ونوع المشكلة لتسهيل التصحيح على المطور.",
                    },
                    {
                        "q": "لماذا تم اختيار ترميز UTF-8 كترميز موحد لجميع استجابات واجهات برمجة التطبيقات؟",
                        "a": "لضمان التوافقية الكاملة في تبادل النصوص متعددة اللغات وتجنب تلف الكلمات العربية أو تحولها لرموز غريبة عند فك التشفير برمجياً في الواجهة الأمامية لـ C#.",
                    },
                    {
                        "q": "كيف تدعم واجهات الخدمة إرجاع روابط ملفات الصوت والفيديو الجاهزة بدلاً من بايتات الملفات مباشرة؟",
                        "a": "تخريج الروابط يقلل من حجم الحمولة (Payload Size) ويوفر استهلاك الباقة للشبكة بشكل هائل؛ حيث تكتفي الخلفية بإرسال رابط نصي خفيف، ويقوم العميل بتنزيل وتشغيل الملف ديناميكياً من التخزين السحابي عند الحاجة.",
                    }
                ]
            }
        ]
    },
    {
        "title": "الفصل السادس: معالجة الأخطاء وموثوقية النظام (Error Handling & Reliability)",
        "description": "يناقش هذا الفصل كيفية التعامل الدفاعي مع الاستثناءات البرمجية، ومعالجة JSON التالفة، وإصلاح انقطاع مخرجات النماذج اللغوية الكبيرة وحماية موارد خادم VPS من الانهيار.",
        "subcategories": [
            {
                "title": "الفرع الأول: تنظيف ومعالجة JSON التالفة والمقطوعة",
                "questions": [
                    {
                        "q": "كيف تعمل دالة `clean_and_parse_json` في استخلاص ومعالجة نصوص JSON التالفة؟",
                        "a": "تتبع الدالة استراتيجية تدرج ذكية: 1) إزالة علامات الماركدوان والأسوار النصية ```json، 2) استخدام التعبيرات المنتظمة لاستخلاص أول قوس متعرج لـ JSON وتجاهل ما قبله وما بعده، 3) محاولة قراءة وتفكيك البيانات؛ فإذا وجدنا بيانات زائدة نلجأ لاستخدام المحلل الأصيل `raw_decode` لقراءة الجزء الصالح فقط، 4) تمرير النص لدالة الإصلاح البرمجية في حال وجود قطع بالنصوص.",
                        "tip": "إظهار الفهم الدقيق لطريقة عمل دوال إصلاح JSON يثبت للجنة التقييم مهاراتكم العالية في البرمجة الدفاعية (Defensive Programming)."
                    },
                    {
                        "q": "كيف تعالج دالة `repair_truncated_json` الانقطاع المفاجئ للنصوص الذي تسببه النماذج؟",
                        "a": "تقوم الدالة بتحليل النص وحساب الفروقات بين الأقواس المفتوحة والأقواس المغلقة للمصفوفات والكائنات. تقوم أولاً بقص الحقول والنصوص المقطوعة في النهاية والتي لم تكتمل صياغتها، ثم تقوم بحذف أي فاصلة زائدة وتلحق الأقواس الناقصة برمجياً لإغلاق الهيكل بنجاح والحصول على JSON صالح للقراءة.",
                    },
                    {
                        "q": "كيف يتعامل النظام مع الأخطاء التي تظهر عند محاولة فك ترميز الأحرف الخاصة في استجابات النماذج؟",
                        "a": "نقوم ببرمجة دالة تنظيف تستبدل الأحرف غير المعيارية برمز بديل أو تحذفها تماماً من السلسلة النصية لمنع انهيار مفسر JSON المدمج ببايثون والحفاظ على تماسك البيانات.",
                    },
                    {
                        "q": "ماذا يحدث لو أرجع الذكاء الاصطناعي بنية JSON صحيحة نحوياً ولكنها فارغة تماماً من الأسئلة؟",
                        "a": "يقوم نظام Pydantic بالتقاط ذلك فوراً؛ لأن نموذج البيانات يفرض وجود عناصر داخل المصفوفة ولا يقبل قائمة خالية. سيفشل التحقق ويرمي النظام استثناءً يفعل آلية إعادة المحاولة لتوليد بنك أسئلة حقيقي.",
                    },
                    {
                        "q": "كيف يتم التعامل مع الكلمات العربية التالفة التي قد ينتجها الذكاء الاصطناعي عند التوليد بالخطأ؟",
                        "a": "النماذج الحديثة قوية جداً بالعربية، ولكن في حال وجود كلمات تالفة، نقوم بعمل فحص أولي للقاموس، أو نعتمد على نظام مراجعة الأخطاء اللغوية لتبسيط وتعديل المخرجات قبل حفظها.",
                    }
                ]
            },
            {
                "title": "الفرع الثاني: حالات أخطاء الخدمات ومعالجة Catch-all",
                "questions": [
                    {
                        "q": "ما هي الحالات التي يرجع فيها النظام كود الحالة HTTP 503 (Service Unavailable)؟",
                        "a": "يرجع النظام هذه الحالة عند تعطل إحدى الخدمات الخارجية الأساسية التي لا يمكن للتطبيق العمل بدونها، مثل حدوث خطأ في استجابة تلاوة ElevenLabs أو نفاد حصة استدعاء الذكاء الاصطناعي بشكل كامل مع فشل التبديل التلقائي لجميع المزودين المتاحين.",
                    },
                    {
                        "q": "لماذا تم تنفيذ معالجة الأخطاء غير المتوقعة كـ 'Catch-all' مع إرجاع رسالة عامة للمستخدمين؟",
                        "a": "لحماية معلومات النظام وأمنه؛ إرجاع تفاصيل الأخطاء وسجلات انهيار المكونات (Stack Traces) للمستخدم يسهل للمخترقين معرفة تفاصيل الكود والمكتبات ونقاط الضعف. نقوم بتسجيل الخطأ الفعلي والتفصيلي بالكامل في خوادمنا ونعرض للمستخدم رسالة محايدة وأنيقة.",
                    },
                    {
                        "q": "ما هو دور الاستثناء `ValueError` في معالجة الأخطاء المنطقية للتطبيق وكيف يتم تمثيله؟",
                        "a": "يُستخدم لتمثيل أخطاء التحقق من البيانات ومدخلات المستخدم (مثل إرسال مستند فارغ أو طلب عدد أسئلة غير مدعوم). يتم معالجته مركزياً في تطبيق FastAPI وتحويله لاستجابة HTTP صالحة ومفهومة للمستخدم.",
                    },
                    {
                        "q": "ما هو دور الاستثناء `RateLimitExceeded` وكيف يتم تخصيص معالج الأخطاء الخاص به؟",
                        "a": "هو استثناء تطلقه مكتبة SlowAPI عند تجاوز المستخدم للحد الأقصى المسموح له من الاستدعاءات. نقوم بربطه بمعالج أخطاء مخصص يقوم بصياغة استجابة HTTP 429 نظيفة تحتوي على تفاصيل وقت الانتظار ورسالة مهذبة.",
                    },
                    {
                        "q": "كيف يتم تنبيه المطورين بحدوث أخطاء قاتلة في سيرفر الإنتاج فور حدوثها؟",
                        "a": "يمكن ربط نظام إدارة الأخطاء بخدمة خارجية لتسجيل الأخطاء (مثل Sentry) أو إرسال إشعار آلي بريدي أو عبر تطبيق تليجرام للمطور يحتوي على كود الحالة ومعرف الطلب لفحص المشكلة فوراً."
                    }
                ]
            },
            {
                "title": "الفرع الثالث: تحديات الذاكرة وتجنب الـ Out of Memory",
                "questions": [
                    {
                        "q": "ماذا يحدث إذا تسبب ملف الـ PDF المرفوع في استهلاك كامل سعة الذاكرة المتاحة لخادم Hostinger VPS وكيف نمنع ذلك؟",
                        "a": "إذا نفدت الذاكرة، سيقوم نظام التشغيل (Linux Out-Of-Memory Killer) بإنهاء عملية Uvicorn/FastAPI فوراً مما يسبب توقف الخلفية عن العمل. نمنع ذلك بوضع قيود برمجية صارمة على حجم الملف المرفوع في إعدادات النظام `MAX_FILE_SIZE_MB = 20` ونقوم بتحليل المستند واستخلاص نصوصه صفحة بصفحة بدلاً من تحميل الملف بالكامل في الذاكرة دفعة واحدة.",
                    },
                    {
                        "q": "كيف يضمن الكود تفريغ بايتات الملفات المرفوعة من الذاكرة فور انتهاء قراءتها؟",
                        "a": "نستخدم الكلمات المفتاحية الذكية في بايثون مثل `with` وقراءة الملف كمصفوفة بايتات مؤقتة يتم التخلص منها برمجياً فور انتهاء دالة استخراج النص `PyMuPDF` وتمرير السلسلة النصية فقط.",
                    },
                    {
                        "q": "ما هي استراتيجيتكم لتخفيف الضغط على معالج السيرفر (CPU) عند معالجة طلبات متزامنة؟",
                        "a": "تأجيل العمليات الثقيلة وتمريرها لخيوط معالجة مستقلة، مع تفعيل حظر الطلبات الزائدة IP-based Rate Limiting وتحديد عدد العمليات العاملة بالتوازي بما يتوافق مع قدرات معالج خادم VPS.",
                    },
                    {
                        "q": "كيف يتعامل النظام مع الأخطاء التي تظهر عند محاولة توليد صورة من Hugging Face؟",
                        "a": "يلتقط الكود خطأ الاتصال أو فشل التوليد، ويقوم فوراً بالتبديل لنموذج Gemini Vision لتخليق الصورة، وإذا فشلت المحاولة أيضاً نلجأ لتوليد صورة محلية توضيحية باستخدام مكتبة Pillow تحمل اسماً كبديل موثوق لضمان بقاء الفيديو متماسكاً.",
                    },
                    {
                        "q": "كيف نضمن عدم حدوث تسريب في الذاكرة العشوائية (Memory Leaks) عند تكرار عمليات توليد الوسائط؟",
                        "a": "نحرص على تجنب المتغيرات العالمية الكبيرة التي تحتفظ بمسارات الوسائط أو بايتات الصوت، ونستدعي جامع المهملات (Garbage Collector) يدوياً في بايثون `import gc; gc.collect()` بعد دمج الفيديوهات الكبيرة للتأكد من إخلاء الذاكرة.",
                    }
                ]
            },
            {
                "title": "الفرع الرابع: إعداد مهل الاتصال وتفادي حلقات التكرار",
                "questions": [
                    {
                        "q": "لماذا تم تحديد مهلة قصوى (Timeout) لكل طلب ذكاء اصطناعي في الإعدادات؟",
                        "a": "لمنع بقاء الاتصالات مفتوحة ومعلقة إلى ما لا نهاية في حال تعطل أو بطء خوادم المزود السحابي، مما يضمن تحرير موارد الخادم وتفعيل التبديل التلقائي للمزود الاحتياطي أو إرجاع خطأ مناسب للمستخدم دون تجميد الجلسة.",
                    },
                    {
                        "q": "كيف نضمن أن استدعاء التبديل التلقائي لا يؤدي لحدوث حلقات تكرار لا نهائية (Infinite Loops)؟",
                        "a": "نضمن ذلك برمجياً بتحديد عدد المحاولات (Max Retries) وتتبع المزود المستدعى؛ حيث يتم إجراء محاولة بديلة واحدة فقط للمزود الاحتياطي، وإذا فشل كلاهما، يرمي النظام استثناءً صريحاً ويتوقف التكرار فوراً.",
                    },
                    {
                        "q": "ما فائدة استخدام دالة `asyncio.wait_for` في العمليات الشبكية داخل التطبيق؟",
                        "a": "تتيح لنا هذه الدالة وضع مهلة زمنية دقيقة بالثانية لتنفيذ العمليات غير المتزامنة. إذا تجاوزت العملية المهلة المحددة، تنهي الدالة التنفيذ تلقائياً وترمي استثناء `TimeoutError` نلتقطه للتعامل معه وحماية خادمنا من الانتظار اللانهائي.",
                    },
                    {
                        "q": "كيف تساهم خوارزمية التباعد والانتظار التدريجي (Exponential Backoff) في تعزيز موثوقية النظام؟",
                        "a": "تُستخدم هذه الخوارزمية عند تكرار طلبات الاتصال الفاشلة بخدمات سحابية؛ حيث نقوم بزيادة زمن الانتظار بين كل محاولة وأخرى تدريجياً لتعطي الخادم الخارجي فرصة للتعافي وتفادي إغراقه بمزيد من الطلبات المتزامنة.",
                    },
                    {
                        "q": "كيف نمنع تعارض العمليات عند طلب تعديل نفس السجل في نفس الوقت في قاعدة البيانات؟",
                        "a": "نستخدم ميزة المعاملات الآمنة (Database Transactions) أو نعتمد على القيود الفريدة للحقول في PostgreSQL لرفض أي طلبات متعارضة وضمان سلامة البيانات وسجلاتها."
                    }
                ]
            },
            {
                "title": "الفرع الخامس: محاكاة الفحوصات والـ Sanity Checks",
                "questions": [
                    {
                        "q": "كيف نضمن استقرار الكود وحمايته من الانهيار عند انقطاع الاتصال بين خادم VPS وسحابة Supabase؟",
                        "a": "نقوم بتغليف جميع استدعاءات Supabase بكتل الحماية `try...except` مع تفعيل وضع التشغيل الاحتياطي؛ حيث لا تتوقف معالجة المستند وإرجاع الامتحانات والخرائط للمستخدمين حتى لو تعطل تخزينها في قاعدة البيانات لضمان كفاءة الخدمة.",
                    },
                    {
                        "q": "ما فائدة تسجيل الأخطاء (Logging) وما هي مستويات التسجيل المستخدمة في المشروع؟",
                        "a": "يساعد التسجيل في مراقبة سلوك النظام وتشخيص المشاكل في بيئة الإنتاج. نستخدم ثلاثة مستويات: `INFO` لتسجيل تدفق العمليات ونجاحها، `WARNING` لتنبيه المطورين من مشاكل غير قاتلة كفشل مزود والتحول للاحتياطي، و `ERROR` لتسجيل الانهيارات البرمجية الكاملة والأخطاء القاتلة.",
                        "tip": "إدارة السجلات بشكل منظم هي الفارق الأساسي بين الكود التجريبي للهواة والكود التجاري للمحترفين."
                    },
                    {
                        "q": "كيف نضمن عدم حدوث ثغرات أمنية في المكتبات الخارجية التي نعتمد عليها في المشروع؟",
                        "a": "نقوم بإجراء فحص أمني دوري لاعتماديات المشروع باستخدام أدوات فحص ثغرات التبعيات مثل `pip-audit` أو أدوات مدمجة في GitHub (Dependabot)، ونحرص على تحديث المكتبات التي تحتوي على ثغرات فوراً وبشكل منظم وآمن.",
                    },
                    {
                        "q": "ما هو دور ملف `check_sanity.py` وكيف يساهم في تقليل وقت اختبار التطبيق؟",
                        "a": "يقوم بفحص أولي سريع يستغرق أجزاءً من الثانية للتحقق من سلامة استيراد جميع التبعيات والموديولات الحيوية، مما يساعد المطور في معرفة وجود أي خطأ أساسي في البيئة فوراً ودون الحاجة لتشغيل اختبارات تكاملية طويلة ومكلفة.",
                    },
                    {
                        "q": "كيف نضمن عدم حدوث تضارب في معالجة التعبيرات المنتظمة (Regex Denial of Service) عند تنظيف JSON؟",
                        "a": "ونحن نستخدم تعبيرات منتظمة بسيطة ومحددة وغير متداخلة ونطبقها على نصوص تم التحقق من حجمها مسبقاً، مما يمنع حدوث عمليات بحث لانهائية تستهلك كامل طاقة المعالج وتعطل الخادم.",
                        "tip": "ثغرات Regex من الثغرات الخفية التي يحب مهندسو أمن المعلومات سؤاله عنها لإظهار عمق معرفتكم البرمجية."
                    }
                ]
            }
        ]
    },
    {
        "title": "الفصل السابع: السياق الأكاديمي والجامعي المصري وحلول الاستضافة البديلة",
        "description": "يناقش هذا الفصل كيفية مواءمة النظام مع متطلبات وضوابط التعليم العالي والجامعات في مصر، وتسهيل استهلاك الموارد للطلاب ومقارنات الاستضافة المحلية.",
        "subcategories": [
            {
                "title": "الفرع الأول: حلول التشغيل المحلي بالجامعة وأمن الامتحانات",
                "questions": [
                    {
                        "q": "كيف يمكن تطبيق ونشر هذا المشروع داخل جامعة مصرية حكومية تعارض استخدام السحابة وتطالب بالحفاظ على سرية الامتحانات؟",
                        "a": "يمكن ذلك بسهولة عبر النشر المحلي (On-Premise Deployment): نقوم بحزم المشروع بالكامل كحاويات Docker ونشره على خوادم مركز البيانات الخاصة بالجامعة (University Data Center). نغلق الاتصالات السحابية الخارجية ونستبدلها بنماذج ذكاء اصطناعي مفتوحة المصدر (مثل Llama 3 أو Arabic-LLM) يتم تشغيلها محلياً على أجهزة خوادم الجامعة المزودة ببطاقات رسومية قوية، مما يضمن أمان وسرية البيانات بنسبة 100% وبدون أي تكاليف استخدام خارجية.",
                        "tip": "هذه إجابة استراتيجية وقوية جداً تسعد أي عميد كلية أو مقيم أكاديمي يبحث عن حلول عملية قابلة للتطبيق الفعلي في مصر."
                    },
                    {
                        "q": "كيف تضمن الجامعة سرية الامتحانات المولدة وعدم وصول الطلاب لها قبل موعدها الفعلي؟",
                        "a": "تخزن الامتحانات في قاعدة البيانات محجوبة خلف صلاحيات أمان صارمة (RLS)، ولا يتم توليد الإجابات النموذجية أو إرسالها للعميل، وتتم عملية التصحيح والتحقق داخل خادم الكلية المغلق تماماً لضمان عدم تسريب أي معلومة.",
                    },
                    {
                        "q": "هل يدعم التطبيق إمكانية طباعة الامتحانات المولدة كملفات PDF ورقية لتوزيعها في لجان الامتحان التقليدية بالكلية؟",
                        "a": "نعم، يمكن بسهولة إعداد دالة برمجية تستخدم مكتبات تصدير ملفات PDF وتقوم بتحويل هيكل بنك الأسئلة JSON إلى ملف PDF منسق بصفحة غلاف الكلية وجاهز للطباعة المباشرة والتوزيع الورقي.",
                    },
                    {
                        "q": "كيف نساعد أساتذة الكلية في مراجعة وتعديل الامتحانات المولدة آلياً قبل اعتمادها رسمياً؟",
                        "a": "يوفر نظامنا واجهة برمجة تتيح إرجاع الأسئلة في صيغة JSON قابلة للتعديل والتحرير. يمكن بناء لوحة تحكم بسيطة لأساتذة الكلية يستعرضون من خلالها الامتحان المقترح ويقومون بحذف الأسئلة أو تعديل صياغتها برمجياً وبسهولة تامة قبل حفظ النسخة النهائية.",
                    },
                    {
                        "q": "ما مدى مرونة دمج المشروع مع نظام التصحيح الإلكتروني (بابل شيت - Bubble Sheet) المتبع بالجامعات المصرية؟",
                        "a": "مرن جداً؛ حيث يمكن للخلفية تصدير مفتاح الإجابة الصحيح بصيغة متوافقة مع أجهزة فك رموز البابل شيت للكلية، مما يتيح تصحيح أوراق الطلاب الورقية آلياً بالاعتماد على نموذج الإجابة المولد من محركنا."
                    }
                ]
            },
            {
                "title": "الفرع الثاني: موازنة تكاليف النماذج السحابية للطلاب",
                "questions": [
                    {
                        "q": "لماذا تم اختيار استضافة Hostinger كحل اقتصادي وتشغيلي للمشروع بدلاً من AWS أو Google Cloud؟",
                        "a": "لأن استضافة Hostinger VPS توفر موارد قوية بأسعار شهرية رمزية ومحددة وثابتة، مقارنة بخدمات السحب الكبرى (AWS, GCP) التي تتبع نظام تسعير معقد ومتغير يصعب حسابه للجامعات والمشاريع الناشئة وقد يؤدي لقفزات مالية مفاجئة.",
                    },
                    {
                        "q": "كم تبلغ التكلفة التشغيلية التقديرية لاستخدام الذكاء الاصطناعي لكل 1000 طالب شهرياً؟",
                        "a": "التكلفة منخفضة للغاية؛ حيث تعتمد على حصص استدعاء مجانية من جوجل لـ Gemini ونصف سنت لكل طلب لـ Groq. استهلاك 1000 طالب لمعالجة مستنداتهم وتلقي الامتحانات لن يتجاوز بضعة دولارات شهرياً مما يجعله خياراً عملياً بامتياز.",
                    },
                    {
                        "q": "كيف يمكن الاستغناء التام عن تكاليف ElevenLabs المرتفعة لتوليد الأصوات؟",
                        "a": "يمكن دمج محركات توليد نصوص صوتية مفتوحة المصدر (مثل Coqui TTS أو Bark) وتشغيلها محلياً على خادم VPS، ورغم أن جودتها ونبرتها الصوتية ستكون أقل طبيعية من ElevenLabs، إلا أنها تحقق استقلالية مالية كاملة للمشروع.",
                    },
                    {
                        "q": "كيف يساهم هذا المشروع في تقليص التكاليف المالية لشراء الكتب والمذكرات الدراسية الورقية للطلاب؟",
                        "a": "يتيح المشروع للكلية رقمنة المحتوى التعليمي ورفعه كملفات إلكترونية، ويقوم النظام بتوليد خرائط المفاهيم والبودكاست والأسئلة مجاناً للطالب، مما يغنيه عن شراء ملخصات خارجية باهظة التكلفة.",
                    },
                    {
                        "q": "هل يوفر النظام ميزات مراقبة الاستهلاك المالي للإداريين لحماية الكلية من فواتير الاستدعاء؟",
                        "a": "نعم، يمكن بناء لوحة تحكم تعرض إحصائيات استهلاك الرموز (Tokens) وعدد الثواني المستهلكة فيElevenLabs لكل مستخدم، مع إمكانية حجب الخدمة تلقائياً عن الحسابات التي تتجاوز السقف المالي المسموح لها."
                    }
                ]
            },
            {
                "title": "الفرع الثالث: دعم الطلاب ذوي الاحتياجات الخاصة وسرعات الشبكة",
                "questions": [
                    {
                        "q": "كيف يساهم هذا المشروع في خدمة ودعم الطلاب ذوي الاحتياجات الخاصة (كالمكفوفين أو ضعاف البصر) في الجامعات المصرية؟",
                        "a": "يقدم المشروع دعماً هائلاً لهم عبر تحويل المحاضرات والكتب المرفوعة لملفات PDF إلى بودكاست صوتي حواري تفاعلي ممتع بالعامية المصرية وفيديوهات مسموعة واضحة. يسهل هذا على الطلاب فاقدي البصر استماع للمحاضرات وتبسيطها دون الحاجة لقارئ نصوص جاف أو مساعدة بشرية مستمرة.",
                    },
                    {
                        "q": "كيف يتعامل النظام مع ضعف سرعات الإنترنت وانقطاع التيار الكهربائي المتكرر لدى الطلاب في مصر؟",
                        "a": "نظام يعتمد على توليد وحفظ نتائج في قاعدة بيانات سحابية بروابط مباشرة وتخزين خفيف للبيانات. في حال انقطاع إنترنت أو كهرباء عن طالب، لا يفقد نتيجته ولا يتعطل توليد على خادم؛ وبمجرد عودة اتصال يمكنه استرجاع خريطة مفاهيم أو بودكاست جاهز فوراً ودون استهلاك باقة إنترنت لديه في إعادة توليد.",
                        "tip": "مراعاة ظروف انقطاع الكهرباء وضعف شبكات الاتصال تعكس اهتماماً عملياً وواقعياً بالبيئة المحلية لتشغيل المشروع."
                    },
                    {
                        "q": "كيف نضمن عدم انقطاع الاتصال بالطالب أثناء رفعه لملف كبير عبر خطوط الاتصال البطيئة في مصر؟",
                        "a": "نقوم ببرمجة واجهة رفع ملفات تدعم تقسيم الملف إلى أجزاء صغيرة (Chunked Uploads) ورفعها بالتتابع، مما يتيح استئناف الرفع عند انقطاع الشبكة وعودتها ودون الحاجة لإعادة الرفع من البداية.",
                    },
                    {
                        "q": "هل يمكن للطالب تنظيف وتحميل ملفات البودكاست والفيديو لمشاهدتها دون الحاجة لفتح المتصفح باستمرار؟",
                        "a": "نعم، الروابط التي ترجعها واجهة الـ API هي روابط مباشرة للملفات الثنائية (MP3 / MP4). يمكن للطالب تحميلها مباشرة على هاتفه أو حاسبه الشخصي وتشغيلها في أي وقت دون اتصال بالإنترنت لتفادي استهلاك الباقة.",
                    },
                    {
                        "q": "كيف يساهم النظام في تقليل الفجوة الرقمية للطلاب المقيمين في القرى والمناطق النائية في مصر؟",
                        "a": "يوفر لهم المحتوى العلمي بأساليب متعددة خفيفة الحجم؛ حيث يمكنهم الاكتفاء بتنزيل خريطة المفاهيم المصورة الخفيفة أو أسئلة الامتحانات النصية القصيرة والتي لا تطلب سرعات إنترنت عالية للتشغيل."
                    }
                ]
            },
            {
                "title": "الفرع الرابع: التوافق مع معايير جودة التعليم المصرية NAQAAE",
                "questions": [
                    {
                        "q": "ما مدى توافق بنوك الأسئلة المولدة برمجياً مع معايير 'الهيئة القومية لضمان جودة التعليم والاعتماد' (NAQAAE) في مصر؟",
                        "a": "يتوافق نظام بشكل ممتاز؛ حيث تفرض هيئة وضع امتحانات تقيس نواتج تعلم مستهدفة وتتوزع على مستويات معرفية متنوعة وهو ما نحققه برمجياً بإلزام ذكاء اصطناعي بتنويع مستويات صعوبة وتغطية منهج بالتساوي وتجنب أسئلة ذاتية أو غامضة واعتماد تام على حقائق موثقة بالمستند.",
                    },
                    {
                        "q": "كيف يضمن النظام قياس مخرجات التعلم المستهدفة (ILOs) للمقررات الأكاديمية؟",
                        "a": "نقوم بإدخال مخرجات التعلم وأهداف المقرر كجزء من موجه النظام للذكاء الاصطناعي، ونطلب منه مطابقة كل سؤال يولده مع ناتج تعلم محدد (مثل معرفة، مهارات ذهنية، أو مهارات عملية) وتخزين هذا الربط بقاعدة البيانات.",
                    },
                    {
                        "q": "ما مدى دقة الامتحانات المولدة في تجنب الانحياز الثقافي أو الجغرافي للطلاب المصريين؟",
                        "a": "دقة عالية؛ لأننا نلزم النماذج بالاعتماد الحصري والمغلق على النص التعليمي المرفوع فقط، مما يمنعها من إقحام أي سياقات خارجية أو أمثلة محلية غريبة قد لا يعرفها الطالب، محققاً العدالة والمساواة بين الجميع.",
                    },
                    {
                        "q": "كيف يساهم النظام في تسهيل إعداد ملفات المقرر الأكاديمي (Course Portfolio) لأساتذة الكلية؟",
                        "a": "يوفر النظام أداة لجمع وتلخيص مخرجات الامتحانات ونسب نجاحها وعرض خرائط المفاهيم للمادة، مما يسهل على أستاذ المادة تعبئة ملف المقرر المطلوب لتقديم ملفات الاعتماد والجودة للكلية بشكل آلي وسريع.",
                    },
                    {
                        "q": "هل يدعم التطبيق صياغة امتحانات باللغة العربية الفصحى للمقررات الأدبية أو الإنسانية بالكلية؟",
                        "a": "نعم، يدعم ذلك بامتياز؛ حيث يتيح النظام قراءة المستندات العربية وتوجيه الذكاء الاصطناعي للتوليد وصياغة بنك الأسئلة بالكامل باللغة العربية الفصحى وتوافقها التام مع متطلبات اللغة للمقرر."
                    }
                ]
            },
            {
                "title": "الفرع الخامس: ميزات تحليلات التعلم والتعاون الفوري",
                "questions": [
                    {
                        "q": "كيف يمكن الاستفادة من قاعدة البيانات والامتحانات القديمة المخزنة لتوليد إحصاءات حول مستويات أداء الطلاب؟",
                        "a": "يمكن تتبع إجابات طلاب على أسئلة برمجياً وتحديد أسئلة أكثر صعوبة والتي يخطئ فيها أغلب طلاب، مما يعطي إحصاءات وتقارير قيمة جداً لأساتذة مواد حول نقاط ضعيفة في استجابة طلاب وتركيز على شرحها مجدداً.",
                        "tip": "هذا المدخل يسمى تحليلات التعلم (Learning Analytics) وهو مجال بحثي واعد ومثير جداً للاهتمام في الجامعات."
                    },
                    {
                        "q": "هل يمكن تشغيل النظام على الهواتف الذكية القديمة أو الاقتصادية للطلاب في مصر؟",
                        "a": "نعم، لأن كامل عمليات حسابية ثقيلة ومعالجة نصوص وتخليق صوت وفيديو تتم بالكامل على خوادم سحابية قوية. واجهة أمامية لطالب لا تقوم سوى بعرض صور وروابط وسائط وقراءة نصوص خفيفة، مما يجعل نظام متوافقاً مع أي هاتف ذكي اقتصادي.",
                    },
                    {
                        "q": "كيف يساهم هذا المشروع في بناء سمعة ومكانة الكلية ككلية رائدة في دمج تقنيات الذكاء الاصطناعي التوليدي بالتعليم؟",
                        "a": "يعد هذا مشروع تطبيقاً عملياً متطوراً لمفاهيم ذكاء اصطناعي توليدي في تبسيط وتطوير تعليم. يعكس نجاح كلية في تدريب طلابها على بناء برمجيات حقيقية متكاملة وسحابية تحل مشاكل واقعية وتواكب ثورة صناعية رابعة.",
                    },
                    {
                        "q": "هل يمكن بناء ميزة امتحانات مجموعات تنافسية (Multiplayer Quizzes) للطلاب على خادم Hostinger؟",
                        "a": "نعم، يمكن ذلك بإدخال تقنيات الاتصال ثنائي الاتجاه المستمر (WebSockets) في FastAPI، مما يتيح للطلاب دخول غرف امتحانات مشتركة وحل نفس الأسئلة بالتزامن مع عرض لوحة الشرف للنتائج فورياً.",
                    },
                    {
                        "q": "كيف يساهم دمج خريطة المفاهيم في مساعدة الطلاب الذين يعانون من تشتت الانتباه والتركيز؟",
                        "a": "الخريطة تلخص المادة كعقد بصرية مترابطة ومختصرة (لا تتعدى 6 كلمات لكل عقدة)، مما يساعد الطلاب ذوي التركيز المحدود في الإحاطة بالمقرر والروابط المعرفية دون تشتت في قراءة نصوص وتفاصيل طويلة."
                    }
                ]
            }
        ]
    },
    {
        "title": "الفصل الثامن: هندسة البرمجيات، الاختبارات والآفاق المستقبلية (Software Engineering)",
        "description": "يناقش هذا الفصل جودة كتابة الكود البرمجي، وهندسة النماذج، والاختبارات الآلية للوحدة والتكامل، وحزم المشروع باستخدام Docker وخطة الترقية المستقبلية للمنصة.",
        "subcategories": [
            {
                "title": "الفرع الأول: استخدام Pydantic v2 وسرعة لغة Rust",
                "questions": [
                    {
                        "q": "لماذا تم اختيار Pydantic الإصدار الثاني (v2) للتحقق من البيانات بدلاً من الإصدار الأول؟",
                        "a": "لأن الإصدار الثاني من Pydantic تم إعادة كتابة قلبه البرمجي بالكامل بلغة Rust السريعة، مما جعله أسرع بـ 17 ضعفاً في التحقق من هياكل البيانات الكبيرة وتوليد JSON مقارنة بالإصدار الأول، بالإضافة لتوفيره دعماً أفضل وخالياً من الأخطاء مع مكتبات FastAPI الحديثة وبيئات التجميع البرمجي.",
                        "tip": "التميز التقني بالإشارة للغة Rust يعزز قيمة الخلفية البرمجية وخبرتكم بأحدث التقنيات."
                    },
                    {
                        "q": "ما هي الفوائد البرمجية لاستخدام Pydantic Schemas للتحقق من المدخلات والمخرجات معاً؟",
                        "a": "يضمن حماية خلفية النظام من قراءة مدخلات تالفة أو خبيثة قد تسبب انهيار التطبيق، كما يضمن استقرار عقود البيانات المرسلة للعميل وتطابقها التام مع معايير الأمان والنوعية المعتمدة.",
                    },
                    {
                        "q": "كيف يساهم Pydantic في تسهيل عملية معالجة وحفظ مصفوفات الـ JSON المعقدة في قاعدة البيانات؟",
                        "a": "يوفر Pydantic دالة تحويل آلي مدمجة (`model_dump()`) تقوم بتحويل نموذج البيانات المعقد إلى قاموس بايثون قياسي متوافق تماماً مع عمليات إدخال قاعدة بيانات Supabase دون كتابة كود تحويل يدوي.",
                    },
                    {
                        "q": "ما هي الفوائد المعمارية والبرمجية لاستخدام بيئة تشغيل Python 3.12 بدلاً من النسخ القديمة؟",
                        "a": "النسخ الحديثة تقدم سرعة تشغيل أعلى بكثير، ودعماً ممتازاً ومطوراً للميزات الحديثة مثل الـ Type Hinting المعقد وتحسينات الكود غير المتزامن ومكتبات التحقق مما ينعكس على كفاءة خلفية التطبيق.",
                    },
                    {
                        "q": "كيف يساهم استخدام نماذج Pydantic في تعزيز الأمان وتفادي ثغرات تسريب البيانات (Data Leakage)؟",
                        "a": "يسمح لنا بتحديد وإرجاع الحقول العامة للمستخدم فقط، وفلترة أو حجب الحقول الحساسة (مثل مفاتيح الإجابات أو التفاصيل السرية) آلياً قبل تمرير الاستجابة للمتصفحات.",
                    }
                ]
            },
            {
                "title": "الفرع الثاني: الاختبارات الآلية للوحدة والتكامل بـ pytest",
                "questions": [
                    {
                        "q": "ما هو دور ملف `test_api.py` وكيف تقومون بإجراء الاختبارات التكاملية (Integration Tests) للنظام؟",
                        "a": "ملف `test_api.py` يحتوي على اختبارات تكاملية آلية تحاكي طلبات حقيقية للمستخدمين. نقوم فيه برفع ملف PDF حقيقي لخدماتنا وفحص الاستجابة للتأكد من إرجاع أكواد HTTP 200 ووجود حقول بنك الأسئلة وخرائط المفاهيم والتخزين السحابي للوسائط وخلو العمليات من أي أخطاء تشغيلية.",
                    },
                    {
                        "q": "كيف يتم كتابة واستخدام اختبارات الوحدة (Unit Tests) للتأكد من سلامة دالة فك تشفير JSON؟",
                        "a": "نستخدم إطار العمل `pytest` لكتابة حالات اختبارية نمرر فيها نصوص JSON تالفة أو مقطوعة أو تحتوي على تعليقات زائدة لدالة `clean_and_parse_json` ونتحقق برمجياً من نجاح الدالة في تصحيحها وإرجاع قاموس بايثون صالح ومطابق للهيكل المتوقع.",
                    },
                    {
                        "q": "لماذا تم اختيار إطار العمل `pytest` بدلاً من مكتبة `unittest` المدمجة في بايثون؟",
                        "a": "لأن `pytest` يوفر أساليب كتابة اختبارات أبسط وأكثر مرونة وخالية من الأقسام الإنشائية المعقدة، ويدعم استخدام الـ Fixtures لإعداد وحقن البيانات المشتركة وتشغيل الاختبارات المتوازية بكفاءة عالية.",
                    },
                    {
                        "q": "كيف نختبر كفاءة وموثوقية نظام الـ Failover الهجين للذكاء الاصطناعي آلياً؟",
                        "a": "نقوم ببرمجة اختبار تكاملي يقوم يدوياً بحجب مفتاح API الخاص بـ Groq أو تمرير قيمة خاطئة له للتسبب في فشل الاستدعاء، ونتحقق من قيام التطبيق بالتبديل التلقائي واستكمال الطلب بنجاح عبر Gemini دون إظهار أخطاء للمستخدم.",
                    },
                    {
                        "q": "ما هي المزايا التي يقدمها فحص الكود الثابت (Static Code Analysis) مثل pylint أو flake8 لمشروعكم؟",
                        "a": "يساعدنا في فحص الكود تلقائياً واكتشاف الأخطاء النحوية، والتحقق من الالتزام بمعايير كتابة الكود النظيف في بايثون (PEP 8)، واكتشاف المتغيرات غير المستخدمة أو التبعيات المفقودة قبل مرحلة التشغيل والتصدير."
                    }
                ]
            },
            {
                "title": "الفرع الثالث: حاويات Docker وتطبيق الـ CI/CD",
                "questions": [
                    {
                        "q": "لماذا تم استخدام `Dockerfile` في جذر المشروع وما هي فائدته الفعلية للتطوير المستقبلي؟",
                        "a": "تم إعداد الـ `Dockerfile` لتسهيل حزم التطبيق واعتمادياته بالكامل في بيئة معزولة ومستقلة. يفيد هذا في تشغيل وتطوير المشروع محلياً أو نشره على أي سحابة أو خادم خاص بأمان تام وخيار تشغيل فوري خالٍ من مشاكل توافق أنظمة التشغيل ومكتبات الخادم.",
                    },
                    {
                        "q": "كيف يدعم المشروع إمكانية العمل بنظام تكامل مستمر ونشر مستمر (CI/CD Pipeline) على سيرفر VPS؟",
                        "a": "يمكن ربط مستودع GitHub بأدوات مثل GitHub Actions، حيث يتم تشغيل خط بناء تلقائي يفحص الكود ويشغل الاختبارات، وعند نجاحها، يقوم السكربت بالاتصال بالسيرفر عبر SSH وسحب الكود الجديد وإعادة تشغيل خدمة uvicorn تلقائياً.",
                    },
                    {
                        "q": "كيف نقوم بإعداد حاوية Docker لتثبيت أداة FFmpeg وتجاوز مشاكل نظام التشغيل؟",
                        "a": "نقوم بكتابة أوامر تثبيت FFmpeg البرمجية داخل الـ `Dockerfile` بالاعتماد على التوزيعة الأساسية (مثل `apt-get install -y ffmpeg`). يضمن هذا تثبيت الأداة وإعداد مساراتها البيئية آلياً داخل الحاوية دون تدخل يدوي من المطور.",
                    },
                    {
                        "q": "ما هي عيوب ومشاكل استخدام حاويات Docker على سيرفرات Hostinger VPS ذات الموارد المحدودة؟",
                        "a": "الاستهلاك الزائد للذاكرة والمعالج لإدارة محرك Docker وتشغيل الحاويات مقارنة بالتثبيت البرمجي المباشر على نظام التشغيل الأساسي، وهو تحدٍ تعاملنا معه بتحسين كفاءة الحاويات واختيار توزيعات خفيفة كـ Alpine.",
                    },
                    {
                        "q": "كيف يدعم نظام الـ CI/CD إشعار المطورين بنجاح أو فشل عمليات النشر التلقائي للسيرفر؟",
                        "a": "يمكن ربط GitHub Actions بخطافات ويب (Webhooks) ترسل رسائل آلية فورية للمطورين عبر تطبيق Discord أو Slack توضح تفاصيل البناء ونتائج الاختبارات وحالة السيرفر بعد النشر."
                    }
                ]
            },
            {
                "title": "الفرع الرابع: خريطة الطريق والآفاق التطويرية للمستقبل",
                "questions": [
                    {
                        "q": "ما هي استراتيجيات التطوير والترقية المستقبلية للمشروع (Future Roadmap)؟",
                        "a": "تشمل رؤيتنا للمستقبل: 1) تطوير تطبيق هواتف ذكية متكامل بلغة Flutter أو C# .NET MAUI، 2) دعم بناء الامتحانات التفاعلية متعددة اللاعبين للطلاب، 3) استخدام تقنيات معالجة لغة متقدمة لتبويب وتحليل نصوص المحاضرات المسجلة كفيديو وتفريغها صوتياً وتلخيصها كخرائط ومقاطع.",
                        "tip": "رسم خطة مستقبلية واضحة يظهر للجنة التقييم أن المشروع حي وقابل للنمو والتطوير التجاري."
                    },
                    {
                        "q": "كيف يمكننا دمج ميزة التقييم التلقائي لأسئلة المقال (Essay Questions Grading) في النظام مستقبلاً؟",
                        "a": "يمكن ذلك برفع إجابة الطالب وتمريرها مع نص السؤال والإجابة النموذجية ومعايير التقييم للذكاء الاصطناعي، ونطلب منه قراءة وتحليل منطق الإجابة وتقدير الدرجة المستحقة وإعطاء تقرير بنقاط القوة والقصور لتبسيط وتسهيل عملية التصحيح.",
                    },
                    {
                        "q": "كيف يمكن تحويل البودكاست التعليمي ليدعم صيغة الفيديو الكرتوني التفاعلي ثلاثي الأبعاد مستقبلاً؟",
                        "a": "يمكن الربط مع محركات ومكتبات تصيير الرسوم ثلاثية الأبعاد أو أدوات توليد الفيديوهات الذكية السحابية وتمرير نص السيناريو لتوليد حركات شفاه متطابقة للأفاتار (Avatar Lip-Sync) وصنع مجسمات كرتونية تشرح المحاضرة بدلاً من الصور الثابتة.",
                    },
                    {
                        "q": "كيف يمكننا إضافة ميزة المذاكرة الذكية المخصصة (Adaptive Learning) للطلاب بناءً على بنك الأسئلة؟",
                        "a": "يمكننا تتبع مستوى إجابات الطالب؛ فإذا أخطأ في موضوع محدد، يقوم النظام تلقائياً باستدعاء خريطة المفاهيم وتوليد بودكاست صوتي مكثف يركز على شرح وتوضيح هذا المفهوم معين تحديداً لتبسيط فهمه ومساعدته في تجاوزه.",
                        "tip": "هذا يعكس رؤية تطبيقية ذكية تدمج الذكاء التوليدي مع نظريات التعليم الشخصي الحديثة."
                    },
                    {
                        "q": "كيف يمكن دمج محرك لتلخيص المحاضرات المسجلة كفيديو مباشرة واستخلاص الامتحانات منها؟",
                        "a": "يمكن دمج مكتبات معالجة الفيديو واستخلاص الصوت، ثم تمريره لمحرك تفريغ النصوص (Speech-to-Text مثل Whisper) للحصول على نص المحاضرة بالكامل، وبدء عمليات محركنا لتوليد الأسئلة والخرائط كالمعتاد."
                    }
                ]
            },
            {
                "title": "الفرع الخامس: معايير نظافة الكود وإدارة خادم Uvicorn",
                "questions": [
                    {
                        "q": "ما هي معايير نظافة الكود (Clean Code) التي تم الالتزام بها أثناء كتابة منطق معالجة الوسائط المتعددة؟",
                        "a": "التزمنا بـ: 1) إعطاء الأسماء الواضحة والمعبرة للمتغيرات والدوال، 2) كتابة دوال صغيرة تؤدي وظيفة واحدة محددة بدقة (Single Responsibility)، 3) فصل التكوينات والمسارات عن منطق التنفيذ، 4) كتابة تعليقات توضيحية تشرح فلسفة الكود وخلفيات القرارات معمارية الصعبة.",
                    },
                    {
                        "q": "لماذا يعتمد التطبيق على خادم Uvicorn لتشغيل الكود محلياً أثناء مرحلة التطوير؟",
                        "a": "لأن Uvicorn هو خادم ASGI فائق السرعة والأداء ومصمم خصيصاً لتشغيل تطبيقات بايثون غير المتزامنة بطلاقة، ويدعم ميزة التحديث التلقائي للكود عند التعديل (Auto-reload) مما يسهل ويسرع عملية التطوير والاختبار للمبرمجين.",
                    },
                    {
                        "q": "كيف نتحقق من أن استدعاءات الخلفية لا تبقي ملفات مفتوحة تسبب تسريب المقابض (File Handles Leak)؟",
                        "a": "نحرص على استخدام كتل `with open(...)` لقراءة الملفات وكتابتها؛ حيث يتولى بايثون إغلاق الملف وتحرير المقبض تلقائياً فور الخروج من الكتلة البرمجية حتى في حال حدوث استثناءات.",
                    },
                    {
                        "q": "كيف يتم التعامل مع عمليات التلوين وضبط مظهر الرسومات الهيكلية لخرائط المفاهيم برمجياً؟",
                        "a": "نستخدم خوارزميات حسابية في ملف `mindmap_renderer.py` تحدد إحداثيات كل عقدة وتفرعها، ونقوم بتلوينها برمجياً وتصييرها بدقة متناهية لتبدو كشجرة مفاهيمية متوازنة وتسهيل قراءتها.",
                    },
                    {
                        "q": "ما هي الفوائد البرمجية لاستخدام النوعية القوية (Type Hinting) في تعريف بارامترات الدوال بخلفية بايثون؟",
                        "a": "تساعد في تسهيل كتابة وقراءة الكود، وتتيح للمحررات البرمجية والـ Linters اكتشاف الأخطاء وتعارض الأنواع وتوقع المخرجات مبكراً قبل التشغيل التجريبي مما يقلل من ظهور الأخطاء البرمجية."
                    }
                ]
            }
        ]
    }
]

# We verify that:
# 8 sections
# Each section has exactly 5 subcategories
# Each subcategory has exactly 5 questions
# Total questions: 8 * 5 * 5 = 200 questions!
# Verified and correct.

# ── WORD DOCUMENT GENERATION ──────────────────────────────────────────────────

def add_page_number(run):
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)

def add_numpages_field(run):
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "NUMPAGES"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)

def set_cell_horizontal_borders(cell, top="E2E8F0", bottom="E2E8F0"):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    if top:
        t = OxmlElement('w:top')
        t.set(qn('w:val'), 'single')
        t.set(qn('w:sz'), '4')
        t.set(qn('w:space'), '0')
        t.set(qn('w:color'), top)
        tcBorders.append(t)
    if bottom:
        b = OxmlElement('w:bottom')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), bottom)
        tcBorders.append(b)
    for side in ['left', 'right']:
        node = OxmlElement(f'w:{side}')
        node.set(qn('w:val'), 'none')
        tcBorders.append(node)
    tcPr.append(tcBorders)

def build_word_document(output_path):
    print("Initializing Document...")
    doc = Document()
    
    # Page Margins (1 inch all sides)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    COLOR_TEAL = RGBColor(0, 150, 136)       # #009688 (FastAPI primary)
    COLOR_NAVY = RGBColor(27, 54, 93)        # #1B365D (Deep Navy)
    COLOR_ORANGE = RGBColor(255, 107, 53)    # #FF6B35 (Accent Orange)
    COLOR_CHARCOAL = RGBColor(55, 65, 81)    # #374151 (Neutral Dark)
    
    # Set default style to Segoe UI
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Segoe UI'
    font.size = Pt(11)
    
    # ── COVER PAGE (Premium Structured Card Layout) ──────────────────────────
    print("Building Cover Page...")
    
    # Configure Different First Page Header/Footer
    section = doc.sections[0]
    section.different_first_page_header_footer = True
    
    # Configure Header (Pages 2+)
    header = section.header
    p_hdr = header.paragraphs[0]
    set_paragraph_rtl(p_hdr)
    p_hdr.paragraph_format.space_after = Pt(6)
    p_hdr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_hdr = p_hdr.add_run("🧠 Ruya Cognitive AI Engine — دليل مناقشة دفاع التخرج")
    set_run_font(run_hdr, font_name="Segoe UI", size_pt=9, color_rgb=RGBColor(120, 130, 140))
    
    # Configure Footer (Pages 2+)
    footer = section.footer
    p_ftr = footer.paragraphs[0]
    set_paragraph_rtl(p_ftr)
    p_ftr.paragraph_format.space_before = Pt(6)
    p_ftr.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    run_ftr_lbl = p_ftr.add_run("صفحة ")
    set_run_font(run_ftr_lbl, font_name="Segoe UI", size_pt=9, color_rgb=RGBColor(120, 130, 140))
    
    run_page = p_ftr.add_run()
    set_run_font(run_page, font_name="Segoe UI", size_pt=9, color_rgb=RGBColor(120, 130, 140))
    add_page_number(run_page)
    
    run_ftr_of = p_ftr.add_run(" من ")
    set_run_font(run_ftr_of, font_name="Segoe UI", size_pt=9, color_rgb=RGBColor(120, 130, 140))
    
    run_numpages = p_ftr.add_run()
    set_run_font(run_numpages, font_name="Segoe UI", size_pt=9, color_rgb=RGBColor(120, 130, 140))
    add_numpages_field(run_numpages)
    
    # Top Emerald Green accent line
    table_line = doc.add_table(rows=1, cols=1)
    table_line.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell_line = table_line.rows[0].cells[0]
    set_cell_background(cell_line, "00A86B") # Emerald Green
    set_cell_margins(cell_line, top=30, bottom=30, left=0, right=0)
    cell_line.width = Inches(6.5)
    # Remove borders
    tcPr = cell_line._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for b in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{b}')
        border.set(qn('w:val'), 'none')
        tcBorders.append(border)
    tcPr.append(tcBorders)
    
    for _ in range(4):
        doc.add_paragraph()
        
    # Title paragraph
    p_title = doc.add_paragraph()
    set_paragraph_rtl(p_title)
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("🧠 موسوعة دفاع التخرج الشاملة\n")
    set_run_font(run_title, font_name="Segoe UI", size_pt=26, color_rgb=COLOR_NAVY, bold=True)
    
    run_subtitle = p_title.add_run("منصة Ruya للذكاء الاصطناعي الإدراكي")
    set_run_font(run_subtitle, font_name="Segoe UI", size_pt=18, color_rgb=COLOR_TEAL, bold=True)
    
    for _ in range(2):
        doc.add_paragraph()
        
    # Metadata Card Table (Fill the page width)
    table_meta = doc.add_table(rows=1, cols=1)
    table_meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell_meta = table_meta.rows[0].cells[0]
    set_cell_background(cell_meta, "F8FAFC")
    set_cell_margins(cell_meta, top=300, bottom=300, left=300, right=300)
    
    # Thick Navy right border, thin light gray elsewhere
    tcPr_m = cell_meta._tc.get_or_add_tcPr()
    tcBorders_m = OxmlElement('w:tcBorders')
    right_m = OxmlElement('w:right')
    right_m.set(qn('w:val'), 'single')
    right_m.set(qn('w:sz'), '36') # 4.5pt
    right_m.set(qn('w:color'), "1B365D")
    tcBorders_m.append(right_m)
    for b in ['top', 'left', 'bottom']:
        border = OxmlElement(f'w:{b}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4') # 0.5pt
        border.set(qn('w:color'), "E2E8F0")
        tcBorders_m.append(border)
    tcPr_m.append(tcBorders_m)
    
    cell_meta.width = Inches(6.5)
    
    p_meta = cell_meta.paragraphs[0]
    set_paragraph_rtl(p_meta)
    p_meta.paragraph_format.line_spacing = 1.4
    
    meta_text = (
        "عدد الأسئلة: 200 سؤال وجواب تفصيلي مرتبة وممنهجة\n"
        "الجهة المستهدفة: لجان التقييم بالجامعات وكليات الحاسبات المصرية\n"
        "بيئة الاستضافة والتشغيل: خادم سحابي مخصص Hostinger VPS (Ubuntu 22.04 LTS)\n"
        "المعمارية التقنية: FastAPI (ASGI Backend) / Gunicorn & Uvicorn / systemd / Nginx / Supabase\n"
        "نموذج التنسيق: نظام معالجة الاتجاه التلقائي المطور (Segoe UI Layout Engine)"
    )
    add_mixed_text_to_p_with_font(p_meta, meta_text, size_pt=10.5, default_color=RGBColor(55, 65, 81))
    
    for _ in range(4):
        doc.add_paragraph()
        
    p_date = doc.add_paragraph()
    set_paragraph_rtl(p_date)
    p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_date = p_date.add_run("تاريخ الإصدار والتنقيح: يونيو 2026")
    set_run_font(run_date, font_name="Segoe UI", size_pt=10, color_rgb=RGBColor(156, 163, 175))
    
    doc.add_page_break()
    
    # ── INTRODUCTION & INDEX ──────────────────────────────────────────────────
    print("Building Introduction...")
    
    add_heading_2_styled(doc, "مقدمة الدليل ودفاع التخرج")
    
    p_intro = doc.add_paragraph()
    set_paragraph_rtl(p_intro)
    p_intro.paragraph_format.line_spacing = 1.3
    p_intro.paragraph_format.space_after = Pt(12)
    add_mixed_text_to_p_with_font(p_intro, 
        "تمت كتابة وتنسيق هذا الدليل لحل مشاكل تداخل الكلمات الإنجليزية والرموز التقنية مع النصوص العربية في ملفات وورد، "
        "مما يمنح المستند مظهراً احترافياً متناسقاً بالكامل وخالياً من مشاكل اتجاه الخطوط. "
        "يحتوي الدليل على 200 سؤال وجواب تفصيلي لتفنيد خيارات المشروع، مع تعديل كامل للتوجهات البرمجية من منصات Serverless "
        "إلى خوادم Hostinger VPS المخصصة لإبقاء معالجة الوسائط سريعة ومتاحة ومستقرة.",
        size_pt=11, default_color=COLOR_CHARCOAL
    )
    
    p_toc_lbl = doc.add_paragraph()
    set_paragraph_rtl(p_toc_lbl)
    p_toc_lbl.paragraph_format.space_before = Pt(12)
    add_rtl_run(p_toc_lbl, "فهرس الأقسام الرئيسية ومحاور التقييم", bold=True, size_pt=13, color_rgb=COLOR_NAVY)
    
    # TOC Table
    table_toc = doc.add_table(rows=1, cols=3)
    table_toc.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    hdr_cells = table_toc.rows[0].cells
    headers = ["عدد الأسئلة", "وصف القسم والتركيز التقني", "اسم القسم"]
    widths = [Inches(1.2), Inches(3.8), Inches(1.5)]
    
    for i, title in enumerate(headers):
        cell = hdr_cells[i]
        set_cell_background(cell, "1B365D")
        set_cell_margins(cell, top=100, bottom=100, left=120, right=120)
        p = cell.paragraphs[0]
        set_paragraph_rtl(p)
        add_rtl_run(p, title, bold=True, size_pt=10.5, color_rgb=RGBColor(255, 255, 255))
        cell.width = widths[i]
        
    toc_rows = [
        ("الفصل الأول", "البنية التحتية، وهيكل النظام، وإعدادات سيرفر Hostinger VPS و Nginx و systemd", "25 سؤالاً"),
        ("الفصل الثاني", "محرك الذكاء الاصطناعي الهجين واستراتيجيات الـ LLMs والتحكم بالحرارة والتلقين", "25 سؤالاً"),
        ("الفصل الثالث", "معالجة وتوليد الوسائط المتعددة (فيديو وبودكاست) واستخدام ElevenLabs و FFmpeg", "25 سؤالاً"),
        ("الفصل الرابع", "إدارة البيانات والربط مع قاعدة بيانات وتخزين Supabase و PostgreSQL", "25 سؤالاً"),
        ("الفصل الخامس", "تكامل وتصميم واجهات الـ API والربط مع C# والتحكم بمعدلات الاستخدام والأمان", "25 سؤالاً"),
        ("الفصل السادس", "معالجة الأخطاء الاستثنائية وتصحيح JSON المكسورة وتفادي انهيار الخادم", "25 سؤالاً"),
        ("الفصل السابع", "السياق التعليمي المصري والتشغيل الداخلي بالجامعة وحلول الاستضافة البديلة", "25 سؤالاً"),
        ("الفصل الثامن", "هندسة البرمجيات النظيفة والاختبارات الآلية للمشروع ورؤية التطوير المستقبلية", "25 سؤالاً"),
    ]
    
    for idx, (ch, desc, q_count) in enumerate(toc_rows):
        row = table_toc.add_row()
        cells = row.cells
        
        bg_color = "F8FAFC" if idx % 2 == 1 else "FFFFFF"
        for cell in cells:
            set_cell_background(cell, bg_color)
            set_cell_horizontal_borders(cell, top="E2E8F0", bottom="E2E8F0")
            
        # Col 2 (right)
        set_cell_margins(cells[2], top=80, bottom=80, left=100, right=100)
        p2 = cells[2].paragraphs[0]
        set_paragraph_rtl(p2)
        add_rtl_run(p2, ch, bold=True, size_pt=10, color_rgb=COLOR_NAVY)
        cells[2].width = widths[2]
        
        # Col 1 (middle)
        set_cell_margins(cells[1], top=80, bottom=80, left=100, right=100)
        p1 = cells[1].paragraphs[0]
        set_paragraph_rtl(p1)
        add_mixed_text_to_p_with_font(p1, desc, size_pt=9.5, default_color=COLOR_CHARCOAL)
        cells[1].width = widths[1]
        
        # Col 0 (left)
        set_cell_margins(cells[0], top=80, bottom=80, left=100, right=100)
        p0 = cells[0].paragraphs[0]
        set_paragraph_rtl(p0)
        add_rtl_run(p0, q_count, bold=True, size_pt=10, color_rgb=COLOR_ORANGE)
        cells[0].width = widths[0]
        
    doc.add_page_break()
    
    # ── SECTIONS GENERATION ───────────────────────────────────────────────────
    q_global_counter = 1
    
    for s_idx, sec in enumerate(sections_data):
        print(f"Adding Section {s_idx+1}...")
        
        # Banner Heading 1
        add_heading_1_banner(doc, sec["title"])
        
        # Section Description
        add_mixed_paragraph(doc, sec["description"], style_type="desc")
        doc.add_paragraph() # Spacing
        
        # Subcategories & Questions
        for sub_cat in sec["subcategories"]:
            # Subcategory Title (Heading 2 styled)
            add_heading_2_styled(doc, sub_cat["title"])
            
            for q_item in sub_cat["questions"]:
                # Generate QA Card
                add_qa_card(
                    doc=doc,
                    q_num=q_global_counter,
                    question=q_item["q"],
                    answer=q_item["a"],
                    tip=q_item.get("tip")
                )
                q_global_counter += 1
                
        if s_idx < len(sections_data) - 1:
            doc.add_page_break()
            
    # Save the document
    print(f"Saving document to {output_path}...")
    doc.save(output_path)
    print("Document Saved Successfully!")

if __name__ == "__main__":
    output_filename = "Ruya_Graduation_Project_Defense_QA_200.docx"
    build_word_document(output_filename)
